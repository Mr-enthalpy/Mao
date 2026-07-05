#!/usr/bin/env python3
"""2月份批量初注脚本"""
import os, shutil
from docx import Document
from lxml import etree

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W14NS = 'http://schemas.microsoft.com/office/word/2010/wordml'
FEB_DIR = r'E:\1952年大传\1952年大传\1952.2'
OUT_PREFIX = '_LLM初注'

LQ = '\u201c'  # "
RQ = '\u201d'  # "

FOOTNOTE_TPL = (
    '<w:footnote xmlns:w="%s" xmlns:w14="%s" w:id="{id}">'
    '<w:p>'
    '<w:pPr>'
    '<w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr>'
    '</w:pPr>'
    '<w:r>'
    '<w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr>'
    '<w:footnoteRef/>'
    '</w:r>'
    '<w:r>'
    '<w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr>'
    '<w:t xml:space="preserve">{text}</w:t>'
    '</w:r>'
    '</w:p>'
    '</w:footnote>'
) % (WNS, W14NS)


def esc(s):
    return s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')


def get_max_footnote_id(fn_root):
    max_id = 0
    for fn in fn_root.findall('{%s}footnote' % WNS):
        fid = int(fn.get('{%s}id' % WNS, '0'))
        max_id = max(max_id, fid)
    return max_id


def insert_footnote_ref(run_elem, fn_id):
    ref_xml = (
        '<w:r xmlns:w="%s">'
        '<w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr>'
        '<w:footnoteReference w:id="%d"/>'
        '</w:r>'
    ) % (WNS, fn_id)
    return etree.fromstring(ref_xml)


def insert_after_text(paragraph, search_text, fn_id):
    full_text = paragraph.text
    if search_text not in full_text:
        return False
    pos_in_full = full_text.index(search_text)
    end_pos = pos_in_full + len(search_text)
    current_pos = 0
    for run in paragraph.runs:
        run_len = len(run.text)
        run_end = current_pos + run_len
        if current_pos <= end_pos <= run_end:
            split_at = end_pos - current_pos
            before = run.text[:split_at]
            after = run.text[split_at:]
            run.text = before
            ref_elem = insert_footnote_ref(run._element, fn_id)
            run._element.addnext(ref_elem)
            if after:
                after_xml = (
                    '<w:r xmlns:w="%s"><w:rPr></w:rPr>'
                    '<w:t xml:space="preserve">%s</w:t></w:r>'
                ) % (WNS, after)
                after_elem = etree.fromstring(after_xml)
                ref_elem.addnext(after_elem)
            return True
        current_pos = run_end
    return False


def add_footnote_to_part(fn_part, fn_id, fn_text):
    fn_root = etree.fromstring(fn_part.blob)
    fn_xml = FOOTNOTE_TPL.format(id=fn_id, text=esc(fn_text))
    fn_elem = etree.fromstring(fn_xml)
    fn_root.append(fn_elem)
    fn_part._blob = etree.tostring(fn_root, xml_declaration=True,
                                     encoding='UTF-8', standalone=True)


def add_end_notes(doc, items):
    if not items:
        return
    sep = doc.add_paragraph()
    sep_run = sep.add_run('\n' + '\u2500' * 40)
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('本篇待人工确认事项：')
    title_run.bold = True
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.add_run('%d. %s' % (i, item))


def process_file(data):
    fname = data['filename']
    src = os.path.join(FEB_DIR, fname)
    out_name = fname.replace('.docx', '_' + OUT_PREFIX + '.docx')
    dst = os.path.join(FEB_DIR, out_name)
    if os.path.exists(dst):
        print('SKIP (exists): ' + out_name)
        return
    shutil.copy2(src, dst)
    doc = Document(dst)

    fn_part = None
    for rel in doc.part.rels.values():
        if 'footnote' in rel.reltype.lower():
            fn_part = rel.target_part
            break
    if fn_part is None:
        print('ERROR: No footnotes part')
        return

    fn_root = etree.fromstring(fn_part.blob)
    next_id = get_max_footnote_id(fn_root) + 1
    print('Processing: ' + fname[:30] + '... (next_id=%d)' % next_id)

    inserted = 0
    skipped = 0
    for search_text, fn_text in data['footnotes']:
        found = False
        for para in doc.paragraphs:
            if search_text in para.text:
                if insert_after_text(para, search_text, next_id):
                    add_footnote_to_part(fn_part, next_id, fn_text)
                    fn_root = etree.fromstring(fn_part.blob)
                    next_id += 1
                    inserted += 1
                    found = True
                    break
        if not found:
            skipped += 1
            print('  NOT FOUND: ' + search_text[:40])

    if data.get('end_notes'):
        add_end_notes(doc, data['end_notes'])

    doc.save(dst)
    print('  -> %d inserted, %d skipped' % (inserted, skipped))


# ═════════════════ ANNOTATION DATA ═════════════════

ANNOTATIONS = {}

# File 01: 与肖劲光等人的谈话
ANNOTATIONS['01'] = {
    'filename': '1952年2月14日，与肖劲光等人的谈话.docx',
    'footnotes': [
        ('建国门', '建国门，北京内城东南侧的城门旧址区域，1950年代在此陆续兴建了外交部等国家部委办公楼。'),
        ('萧劲光', '萧劲光（1903\u20141989），湖南长沙人，1922年加入中国共产党。中国人民解放军海军主要创建者之一，时任海军司令员。1955年被授予大将军衔。'),
        ('王宏坤', '王宏坤（1909\u20141993），湖北麻城人，1929年参加中国工农红军。时任海军副司令员。1955年被授予上将军衔。'),
        ('刘道生', '刘道生（1915\u20141995），湖南茶陵人，1930年参加红军。时任海军副政治委员。1955年被授予中将军衔。'),
        ('罗舜初', '罗舜初（1914\u20141981），福建上杭人，1931年参加红军。时任海军参谋长。1955年被授予中将军衔。'),
        ('罗瑞卿', '罗瑞卿（1906\u20141978），四川南充人，1928年加入中国共产党。时任中央人民政府公安部部长。'),
        ('刘亚楼', '刘亚楼（1910\u20141965），福建武平人，1929年加入中国共产党。时任中国人民解放军空军司令员。1955年被授予上将军衔。'),
        (LQ + '三反' + RQ, LQ + '三反' + RQ + '运动，指1951年底至1952年10月在中国党政机关、国营企业等单位开展的' + LQ + '反贪污、反浪费、反官僚主义' + RQ + '运动。'),
        ('抗美援朝', '抗美援朝，即中国人民志愿军赴朝鲜支援朝鲜人民抗击以美国为首的联合国军的战争。1950年10月开始，至1953年7月签订停战协定。'),
        ('2亿卢布', '卢布，苏联货币单位。新中国成立初期，中苏贸易以卢布结算。此处' + LQ + '2亿卢布' + RQ + '指中苏贸易协定中用于国防采购的外汇额度。'),
        ('江南造船厂', '江南造船厂，位于上海，前身为1865年创办的江南机器制造总局，是中国近代最早的造船企业之一。新中国成立后恢复军用和民用船舶建造。'),
        ('青岛造船厂', '青岛造船厂，位于山东青岛，新中国成立初期承担海军小型舰艇的建造任务。'),
    ],
    'comments': [],
    'end_notes': [
        '【需核】' + LQ + '2亿卢布' + RQ + '的具体数额和对苏联军购的对应关系，建议核对相关外交档案。',
    ],
}

# File 02: 转发华东军区"打虎"批示
ANNOTATIONS['02'] = {
    'filename': '1952年2月20日，在转发华东军区党委关于' + LQ + '打虎' + RQ + '情况和部署的报告上的批示.docx',
    'footnotes': [
        (LQ + '打虎' + RQ, LQ + '打虎' + RQ + '，' + LQ + '三反' + RQ + '运动中的术语，指查处和打击贪污数额巨大的贪污分子（即' + LQ + '大老虎' + RQ + '），是运动的重点工作。'),
        ('华东军区', '华东军区，中国人民解放军一级军区之一，1947年1月由山东军区与华中军区合并成立。司令员陈毅，政治委员饶漱石。辖山东、江苏、安徽、浙江、福建等省军事工作。'),
        ('逼供信', '逼供信，指在审案办案中采用逼供、诱供、指供等非法手段迫使嫌疑人供认并加以采信的错误做法。' + LQ + '三反' + RQ + '运动后期中央反复强调严禁逼供信。'),
    ],
    'comments': [],
    'end_notes': [
        '【需核】华东军区党委关于' + LQ + '打虎' + RQ + '情况的原始报告具体内容，建议核查相关档案。',
    ],
}

# File 03: 三反与整党结合
ANNOTATIONS['03'] = {
    'filename': '1952年2月3日，关于' + LQ + '三反' + RQ + '运动和整党运动结合进行的指示.docx',
    'footnotes': [
        (LQ + '三反' + RQ + '运动', LQ + '三反' + RQ + '运动，指1951年底至1952年10月在中国党政机关、国营企业等单位开展的' + LQ + '反贪污、反浪费、反官僚主义' + RQ + '运动。'),
        ('整党运动', '整党运动，指1951年至1954年间在全党开展的整党运动，分三个阶段进行。主要任务是进行党员标准教育，对全体党员进行登记、审查和处理，清除不合格党员，纯洁党的组织。'),
        ('中南局', '中共中央中南局，中共中央在中南地区的代表机关。1949年成立，驻武汉，辖河南、湖北、湖南、江西、广东、广西六省区党的工作。第一书记林彪。'),
        ('八项标准', '党员八项标准，即1951年3月第一次全国组织工作会议通过的《共产党员标准的八项条件》，是整党运动中衡量党员是否合格的基本依据。'),
        ('中央人民政府各党组', '指中央人民政府各部、委、院、署等机关中设立的中国共产党党组。'),
        ('志愿军', '中国人民志愿军，1950年10月组成并开赴朝鲜参加抗美援朝战争。'),
        ('兵团', '兵团，中国人民解放军在解放战争后期至1950年代设置的一级编制单位，介于野战军（大军区）和军之间。'),
    ],
    'comments': [],
    'end_notes': [
        '【需核】中南局组织部《关于将整党报告改为' + LQ + '三反' + RQ + '报告的通知》的具体内容和发文时间需要核实。',
    ],
}

# File 04: 工矿企业三反
ANNOTATIONS['04'] = {
    'filename': '1952年2月4日，关于工矿企业如何进行' + LQ + '三反' + RQ + '运动的指示.docx',
    'footnotes': [
        (LQ + '三反' + RQ + '运动', LQ + '三反' + RQ + '运动，指1951年底至1952年10月在中国党政机关、国营企业等单位开展的' + LQ + '反贪污、反浪费、反官僚主义' + RQ + '运动。'),
        ('增产节约', '增产节约运动，新中国成立初期在全国范围内开展的以增加生产、厉行节约为中心内容的群众运动。1951年10月毛泽东向全国发出开展增产节约运动的号召。'),
        ('民主改革', '民主改革，指新中国成立初期在工矿企业中进行的社会改革运动，主要内容包括废除封建把头制度、清除反革命分子、建立职工代表会议等民主管理制度。'),
        ('经济核算制', '经济核算制，社会主义国营企业管理的基本制度之一，要求企业以货币形式独立核算生产经营的消耗和成果，以收抵支，保证盈利。'),
    ],
    'comments': [],
    'end_notes': [],
}

# File 05: 与李家骥的谈话
ANNOTATIONS['05'] = {
    'filename': '1952年2月5日、6日，与李家骥的谈话.docx',
    'footnotes': [
        ('李家骥', '李家骥，毛泽东的警卫人员（卫士），负责毛泽东的日常生活起居和安全保卫工作。'),
        ('没有调查就没有发言权', LQ + '没有调查，就没有发言权' + RQ + '，毛泽东在1930年5月《反对本本主义》（原题为《调查工作》）一文中提出的著名论断，强调调查研究对于正确决策的根本重要性。'),
    ],
    'comments': [],
    'end_notes': [],
}

# File 06: 和李敏的谈话
ANNOTATIONS['06'] = {
    'filename': '1952年2月8日，和李敏的谈话.docx',
    'footnotes': [
        ('李敏', '李敏（1936\u2014），毛泽东与贺子珍之女，原名毛姣姣（小名娇娇）。1936年生于陕北保安，1941年被送往苏联，1947年回国。1949年回到毛泽东身边，在北京读书生活。1952年春节为1月27日，中小学寒假一般持续至正月十五前后，2月8日（农历正月十三）确在寒假期间。'),
    ],
    'end_notes': [],
}


# ═════════════════ MAIN ═════════════════

def main():
    all_files = [f for f in os.listdir(FEB_DIR) if f.endswith('.docx') and not f.startswith('__')]

    for key, data in ANNOTATIONS.items():
        fname = data['filename']
        matched = None
        for af in all_files:
            if af.startswith(fname[:15]):
                matched = af
                break
        if matched:
            data['filename'] = matched
            process_file(data)
        else:
            print('FILE NOT FOUND: ' + fname[:30])

    # Clean up test files
    for f in os.listdir(FEB_DIR):
        if f.startswith('__'):
            os.remove(os.path.join(FEB_DIR, f))

    print('\n2月份初注完成!')


if __name__ == '__main__':
    main()
