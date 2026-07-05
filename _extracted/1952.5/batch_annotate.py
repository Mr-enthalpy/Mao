#!/usr/bin/env python3
"""5月份批量初注脚本"""
import os, shutil
from docx import Document
from lxml import etree

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W14NS = 'http://schemas.microsoft.com/office/word/2010/wordml'
MONTH_DIR = r'E:\1952年大传\1952年大传\1952.5'
OUT_PREFIX = '_LLM初注'
LQ = '\u201c'; RQ = '\u201d'

FOOTNOTE_TPL = (
    '<w:footnote xmlns:w="%s" xmlns:w14="%s" w:id="{id}">'
    '<w:p><w:pPr><w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr></w:pPr>'
    '<w:r><w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr><w:footnoteRef/></w:r>'
    '<w:r><w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr><w:t xml:space="preserve">{text}</w:t></w:r>'
    '</w:p></w:footnote>'
) % (WNS, W14NS)


def esc(s):
    return s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')


def get_max_footnote_id(fn_root):
    max_id = 0
    for fn in fn_root.findall('{%s}footnote' % WNS):
        max_id = max(max_id, int(fn.get('{%s}id' % WNS, '0')))
    return max_id


def insert_after_text(paragraph, search_text, fn_id):
    full_text = paragraph.text
    if search_text not in full_text:
        return False
    pos_in_full = full_text.index(search_text)
    end_pos = pos_in_full + len(search_text)
    current_pos = 0
    for run in paragraph.runs:
        run_len = len(run.text)
        run_end = current_pos + run_len
        if current_pos <= end_pos <= run_end:
            split_at = end_pos - current_pos
            after = run.text[split_at:]
            run.text = run.text[:split_at]
            ref_xml = (
                '<w:r xmlns:w="%s">'
                '<w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr>'
                '<w:footnoteReference w:id="%d"/>'
                '</w:r>'
            ) % (WNS, fn_id)
            ref_elem = etree.fromstring(ref_xml)
            run._element.addnext(ref_elem)
            if after:
                after_xml = (
                    '<w:r xmlns:w="%s"><w:rPr></w:rPr>'
                    '<w:t xml:space="preserve">%s</w:t></w:r>'
                ) % (WNS, after)
                ref_elem.addnext(etree.fromstring(after_xml))
            return True
        current_pos = run_end
    return False


def add_footnote_to_part(fn_part, fn_id, fn_text):
    fn_root = etree.fromstring(fn_part.blob)
    fn_xml = FOOTNOTE_TPL.format(id=fn_id, text=esc(fn_text))
    fn_root.append(etree.fromstring(fn_xml))
    fn_part._blob = etree.tostring(fn_root, xml_declaration=True, encoding='UTF-8', standalone=True)


def add_end_notes(doc, items):
    if not items:
        return
    sep = doc.add_paragraph()
    sep.add_run('\n' + '\u2500' * 40)
    title_para = doc.add_paragraph()
    run = title_para.add_run('本篇待人工确认事项：')
    run.bold = True
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.add_run('%d. %s' % (i, item))


def process_file(data):
    fname = data['filename']
    src = os.path.join(MONTH_DIR, fname)
    out_name = fname.replace('.docx', '_' + OUT_PREFIX + '.docx')
    dst = os.path.join(MONTH_DIR, out_name)
    if os.path.exists(dst):
        print('SKIP: ' + out_name)
        return
    shutil.copy2(src, dst)
    doc = Document(dst)
    fn_part = None
    for rel in doc.part.rels.values():
        if 'footnote' in rel.reltype.lower():
            fn_part = rel.target_part
            break
    if fn_part is None:
        print('ERROR: no footnotes')
        return
    fn_root = etree.fromstring(fn_part.blob)
    next_id = get_max_footnote_id(fn_root) + 1
    print('>>> ' + fname[:30] + ' (id=%d)' % next_id)

    inserted = 0
    for search_text, fn_text in data['footnotes']:
        found = False
        for para in doc.paragraphs:
            if search_text in para.text:
                if insert_after_text(para, search_text, next_id):
                    add_footnote_to_part(fn_part, next_id, fn_text)
                    fn_root = etree.fromstring(fn_part.blob)
                    next_id += 1
                    inserted += 1
                    found = True
                    break
        if not found:
            print('  NOT FOUND: ' + search_text[:50])
    if data.get('end_notes'):
        add_end_notes(doc, data['end_notes'])
    doc.save(dst)
    print('  -> %d footnotes' % inserted)


# ═════════════════ ANNOTATIONS ═════════════════
ANNOTATIONS = {}

# --- File 01: 推迟三反五反 ---
S3 = LQ + '三反' + RQ
S5 = LQ + '五反' + RQ
ANNOTATIONS['01'] = {
    'filename': '1952年5月23日，关于推迟县区乡的' + S3 + '和中小城市的' + S5 + '的指示.docx',
    'footnotes': [
        (S3, S3 + '运动，指1951年底至1952年10月在中国党政机关、国营企业等单位开展的' + LQ + '反贪污、反浪费、反官僚主义' + RQ + '运动。'),
        (S5, S5 + '运动，指1952年上半年在中国资本主义工商业者中开展的' + LQ + '反行贿、反偷税漏税、反盗骗国家财产、反偷工减料、反盗窃国家经济情报' + RQ + '的运动。'),
        ('中南局', '中共中央中南局，中共中央在中南地区的代表机关，辖河南、湖北、湖南、江西、广东、广西六省区党的工作。'),
        ('湖北省委', '即中国共产党湖北省委员会，1949年5月武汉解放后成立，李先念任书记（1949-1954）。'),
        ('秋征', '秋征，即秋季征收农业税（公粮），新中国成立初期国家财政收入的重要来源之一。'),
        ('土改复查', '土地改革复查，指土地改革完成后，对土改中遗留问题和偏差进行的检查和纠正工作。'),
        ('广西省委', '即中国共产党广西省委员会，1949年12月广西解放后成立。1958年广西壮族自治区成立后改为自治区党委。'),
    ],
    'end_notes': [
        '【需核】中南局5月15日和湖北省委5月10日关于县区' + S3 + '计划的电报原文内容需要核查。',
        '【需核】广西省委整顿土改工作队的经验具体内容需要核实。',
    ],
}

# --- File 02: 三反基础上整党建党 ---
ANNOTATIONS['02'] = {
    'filename': '1952年5月30日，关于在' + S3 + '运动的基础上进行整党建党工作的指示.docx',
    'footnotes': [
        (S3 + '运动', S3 + '运动，指1951年底至1952年10月在中国党政机关、国营企业等单位开展的' + LQ + '反贪污、反浪费、反官僚主义' + RQ + '运动。'),
        ('8项标准', '党员八项标准，即1951年3月第一次全国组织工作会议通过的《共产党员标准的八项条件》，是整党运动中衡量党员是否合格的基本依据。'),
        (S5, S5 + '运动，指1952年上半年在资本主义工商业者中开展的' + LQ + '反行贿、反偷税漏税、反盗骗国家财产、反偷工减料、反盗窃国家经济情报' + RQ + '的运动。'),
        ('民主改革', '民主改革，指新中国成立初期在工矿企业中进行的社会改革运动，主要内容包括废除封建把头制度、建立职工代表会议等民主管理制度。'),
        ('增产节约', '增产节约运动，新中国成立初期在全国范围内开展的以增加生产、厉行节约为中心内容的群众运动。'),
        ('青年团员', '指中国新民主主义青年团团员。新民主主义青年团成立于1949年4月，1957年改名为中国共产主义青年团。'),
        ('抗美援朝', '抗美援朝，即中国人民志愿军赴朝鲜支援朝鲜人民抗击以美国为首的联合国军的战争，1950年10月开始。'),
        ('清理' + LQ + '中层' + RQ, LQ + '清理中层' + RQ + '，指1951年至1952年在知识分子中进行的思想改造运动中的一个阶段，主要清理隐藏在国家机关和文教部门干部队伍中的反革命分子和坏分子。'),
        ('关门主义', '关门主义，指拒绝吸收符合条件的积极分子入党的错误倾向。毛泽东多次批评关门主义，强调要大胆发展党员。'),
        ('民主建政', '民主建政，指新中国成立初期在农村基层建立人民民主政权的工作，包括建立乡（村）人民代表大会和人民政府。'),
        ('全国组织工作会议', '即1951年3月在北京召开的中国共产党第一次全国组织工作会议。刘少奇作报告，会议通过了《关于整顿党的基层组织的决议》和共产党员标准的八项条件。'),
    ],
    'end_notes': [
        '【需核】文中' + LQ + '60万青年团员' + RQ + '和' + LQ + '20%' + RQ + '等发展党员的统计数据是否准确，建议核对。',
        '【需核】' + LQ + '360万产业职工' + RQ + '、' + LQ + '150万店员' + RQ + '、' + LQ + '12万个新区乡村' + RQ + '等统计数据的出处需要核实。',
    ],
}

# --- File 03: 与丁玲的谈话 ---
ANNOTATIONS['03'] = {
    'filename': '1952年5月，与丁玲的谈话.docx',
    'footnotes': [
        ('丁玲', '丁玲（1904\u20141986），原名蒋伟，湖南临澧人，中国现代女作家。代表作有《莎菲女士的日记》《太阳照在桑干河上》。时任中国作家协会副主席、《文艺报》主编。'),
        ('周扬', '周扬（1908\u20141989），湖南益阳人，文艺理论家。1937年到延安，曾任鲁迅艺术学院院长。时任中共中央宣传部副部长、文化部副部长。'),
    ],
    'end_notes': [
        '【需核】该谈话原文前的' + LQ + '（……）' + RQ + '省略号表明原文有删节，删节的具体内容和原因需核实。',
        '【需核】丁玲与周扬在1952年的职务和相互关系（二人后来在文艺界整风中有严重冲突），建议核查相关文学史资料。',
    ],
}


# ═════════════════ MAIN ═════════════════
def main():
    all_files = [f for f in os.listdir(MONTH_DIR) if f.endswith('.docx') and not f.startswith('__')]
    for key, data in ANNOTATIONS.items():
        fname = data['filename']
        matched = None
        for af in all_files:
            if af.startswith(fname[:15]):
                matched = af
                break
        if matched:
            data['filename'] = matched
            process_file(data)
        else:
            print('NOT FOUND: ' + fname[:30])
    for f in os.listdir(MONTH_DIR):
        if f.startswith('__'):
            os.remove(os.path.join(MONTH_DIR, f))
    print('\n5月初注完成!')


if __name__ == '__main__':
    main()
