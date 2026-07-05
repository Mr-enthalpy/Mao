#!/usr/bin/env python3
"""9月份批量初注脚本"""
import os, shutil
from docx import Document
from lxml import etree

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W14NS = 'http://schemas.microsoft.com/office/word/2010/wordml'
MONTH_DIR = r'E:\1952年大传\1952年大传\1952.9'
OUT_PREFIX = '_LLM初注'
LQ='\u201c'; RQ='\u201d'

FOOTNOTE_TPL = (
    '<w:footnote xmlns:w="%s" xmlns:w14="%s" w:id="{id}">'
    '<w:p><w:pPr><w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr></w:pPr>'
    '<w:r><w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr><w:footnoteRef/></w:r>'
    '<w:r><w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr><w:t xml:space="preserve">{text}</w:t></w:r>'
    '</w:p></w:footnote>'
) % (WNS, W14NS)

def esc(s): return s.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')
def get_max_id(fn_root): return max(int(fn.get('{%s}id'%WNS,'0')) for fn in fn_root.findall('{%s}footnote'%WNS))

def insert_after_text(paragraph, search_text, fn_id):
    full_text = paragraph.text
    if search_text not in full_text: return False
    end_pos = full_text.index(search_text) + len(search_text); cur = 0
    for run in paragraph.runs:
        run_end = cur + len(run.text)
        if cur <= end_pos <= run_end:
            sp = end_pos - cur; after = run.text[sp:]; run.text = run.text[:sp]
            ref = etree.fromstring(('<w:r xmlns:w="%s"><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr><w:footnoteReference w:id="%d"/></w:r>')%(WNS,fn_id))
            run._element.addnext(ref)
            if after:
                a = etree.fromstring(('<w:r xmlns:w="%s"><w:rPr></w:rPr><w:t xml:space="preserve">%s</w:t></w:r>')%(WNS,after))
                ref.addnext(a)
            return True
        cur = run_end
    return False

def add_fn(fn_part, fn_id, fn_text):
    r = etree.fromstring(fn_part.blob)
    r.append(etree.fromstring(FOOTNOTE_TPL.format(id=fn_id, text=esc(fn_text))))
    fn_part._blob = etree.tostring(r, xml_declaration=True, encoding='UTF-8', standalone=True)

def add_end(doc, items):
    if not items: return
    doc.add_paragraph().add_run('\n' + '\u2500' * 40)
    r = doc.add_paragraph().add_run('本篇待人工确认事项：'); r.bold = True
    for i, item in enumerate(items, 1): doc.add_paragraph().add_run('%d. %s' % (i, item))

def process_file(data):
    fname = data['filename']; src = os.path.join(MONTH_DIR, fname)
    dst = os.path.join(MONTH_DIR, fname.replace('.docx','_'+OUT_PREFIX+'.docx'))
    if os.path.exists(dst): print('SKIP'); return
    shutil.copy2(src, dst); doc = Document(dst)
    fn_part = None
    for rel in doc.part.rels.values():
        if 'footnote' in rel.reltype.lower(): fn_part = rel.target_part; break
    if fn_part is None: print('ERROR'); return
    next_id = get_max_id(etree.fromstring(fn_part.blob)) + 1
    print('>>> ' + fname[:30] + ' (id=%d)' % next_id)
    ins = 0
    for st, ft in data['footnotes']:
        for para in doc.paragraphs:
            if st in para.text:
                if insert_after_text(para, st, next_id): add_fn(fn_part, next_id, ft); next_id += 1; ins += 1; break
        else: print('  NF: ' + st[:50])
    if data.get('end_notes'): add_end(doc, data['end_notes'])
    doc.save(dst); print('  -> %d' % ins)

S3=LQ+'三反'+RQ

ANNOTATIONS = {
    '01': {
        'filename': '1952年9月1日，关于培养高等、中等学校马克思列宁主义理论师资的指示.docx',
        'footnotes': [
            (S3 + '运动', S3 + '运动，指1951年底至1952年10月在中国党政机关、国营企业等单位开展的' + LQ + '反贪污、反浪费、反官僚主义' + RQ + '运动。'),
            ('组织清理', LQ + '组织清理' + RQ + '，指在知识分子思想改造运动中对隐藏在国家机关和文教部门中的反革命分子和坏分子进行组织上的审查和清理。此项工作与' + LQ + '清理中层' + RQ + '密切相关。'),
            ('中国人民大学', '中国人民大学，1950年10月3日在北京成立，以华北大学为基础组建，吴玉章任校长。是新中国成立后创办的第一所新型正规大学，以培养财经、政法和马克思主义理论干部为主要任务。'),
            ('中央教育部', '中央人民政府教育部，1949年10月成立，马叙伦任部长，主管全国教育工作。1952年11月分设为高等教育部和教育部。'),
            ('华北局', '中共中央华北局，中共中央在华北地区的代表机关，驻北京。辖北京市、天津市、河北省、山西省和原绥远省（1954年撤销）。'),
            ('华北五省二市', '指河北省、山西省、平原省（1952年11月撤销）、绥远省（1954年撤销）、察哈尔省（1952年11月撤销）及北京、天津两市。这是1952年9月时的华北行政区划。'),
        ],
        'end_notes': [
            '【需核】各中央局拟订的政治理论师资培养计划是否在9月20日前上报中央，以及中国人民大学马克思列宁主义研究班第一期是否按期开学，建议核实。',
        ],
    },
}

def main():
    all_files = [f for f in os.listdir(MONTH_DIR) if f.endswith('.docx') and not f.startswith('__')]
    for key, data in ANNOTATIONS.items():
        fname = data['filename']; matched = None
        for af in all_files:
            if af.startswith(fname[:15]): matched = af; break
        if matched: data['filename'] = matched; process_file(data)
        else: print('NF: ' + fname[:30])
    for f in os.listdir(MONTH_DIR):
        if f.startswith('__'): os.remove(os.path.join(MONTH_DIR, f))
    print('\n9月初注完成!')

if __name__ == '__main__': main()
