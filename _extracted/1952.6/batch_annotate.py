#!/usr/bin/env python3
"""6月份批量初注脚本"""
import os, shutil
from docx import Document
from lxml import etree

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W14NS = 'http://schemas.microsoft.com/office/word/2010/wordml'
MONTH_DIR = r'E:\1952年大传\1952年大传\1952.6'
OUT_PREFIX = '_LLM初注'
LQ='\u201c'; RQ='\u201d'; ANG='\u300a'; ANR='\u300b'

FOOTNOTE_TPL = (
    '<w:footnote xmlns:w="%s" xmlns:w14="%s" w:id="{id}">'
    '<w:p><w:pPr><w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr></w:pPr>'
    '<w:r><w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr><w:footnoteRef/></w:r>'
    '<w:r><w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr><w:t xml:space="preserve">{text}</w:t></w:r>'
    '</w:p></w:footnote>'
) % (WNS, W14NS)

def esc(s):
    return s.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')

def get_max_footnote_id(fn_root):
    return max(int(fn.get('{%s}id' % WNS, '0')) for fn in fn_root.findall('{%s}footnote' % WNS))

def insert_after_text(paragraph, search_text, fn_id):
    full_text = paragraph.text
    if search_text not in full_text: return False
    pos_in_full = full_text.index(search_text)
    end_pos = pos_in_full + len(search_text)
    current_pos = 0
    for run in paragraph.runs:
        run_end = current_pos + len(run.text)
        if current_pos <= end_pos <= run_end:
            split_at = end_pos - current_pos
            after = run.text[split_at:]
            run.text = run.text[:split_at]
            ref_xml = ('<w:r xmlns:w="%s"><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr><w:footnoteReference w:id="%d"/></w:r>') % (WNS, fn_id)
            ref_elem = etree.fromstring(ref_xml)
            run._element.addnext(ref_elem)
            if after:
                after_xml = ('<w:r xmlns:w="%s"><w:rPr></w:rPr><w:t xml:space="preserve">%s</w:t></w:r>') % (WNS, after)
                ref_elem.addnext(etree.fromstring(after_xml))
            return True
        current_pos = run_end
    return False

def add_footnote(fn_part, fn_id, fn_text):
    fn_root = etree.fromstring(fn_part.blob)
    fn_root.append(etree.fromstring(FOOTNOTE_TPL.format(id=fn_id, text=esc(fn_text))))
    fn_part._blob = etree.tostring(fn_root, xml_declaration=True, encoding='UTF-8', standalone=True)

def add_end_notes(doc, items):
    if not items: return
    doc.add_paragraph().add_run('\n' + '\u2500' * 40)
    r = doc.add_paragraph().add_run('本篇待人工确认事项：'); r.bold = True
    for i, item in enumerate(items, 1):
        doc.add_paragraph().add_run('%d. %s' % (i, item))

def process_file(data):
    fname = data['filename']
    src = os.path.join(MONTH_DIR, fname)
    out_name = fname.replace('.docx', '_' + OUT_PREFIX + '.docx')
    dst = os.path.join(MONTH_DIR, out_name)
    if os.path.exists(dst): print('SKIP'); return
    shutil.copy2(src, dst)
    doc = Document(dst)
    fn_part = None
    for rel in doc.part.rels.values():
        if 'footnote' in rel.reltype.lower(): fn_part = rel.target_part; break
    if fn_part is None: print('ERROR: no fn'); return
    next_id = get_max_footnote_id(etree.fromstring(fn_part.blob)) + 1
    print('>>> ' + fname[:30] + ' (id=%d)' % next_id)
    inserted = 0
    for st, ft in data['footnotes']:
        for para in doc.paragraphs:
            if st in para.text:
                if insert_after_text(para, st, next_id):
                    add_footnote(fn_part, next_id, ft); next_id += 1; inserted += 1; break
        else:
            print('  NOT FOUND: ' + st[:50])
    if data.get('end_notes'): add_end_notes(doc, data['end_notes'])
    doc.save(dst)
    print('  -> %d' % inserted)

# ═════════════════ ANNOTATIONS ═════════════════
ANNOTATIONS = {}
S3 = LQ + '三反' + RQ; S5 = LQ + '五反' + RQ

# 01: 与周恩来的谈话
ANNOTATIONS['01'] = {
    'filename': '1952年6月13日，与周恩来的谈话.docx',
    'footnotes': [
        ('完全守法户', '完全守法户，' + S5 + '运动中对私营工商户的五级分类中的第一类，指没有违法行为、完全遵守国家法令的工商户。五级分类依次为：守法户、基本守法户、半守法半违法户、严重违法户、完全违法户。'),
    ],
    'end_notes': [
        '【需核】本谈话的完整上下文（毛泽东是在什么场合、就哪家工商户的分类问题说这番话的）需要核查。',
    ],
}

# 02: 三反运动结束指示
ANNOTATIONS['02'] = {
    'filename': '1952年6月15日，中共中央关于争取胜利结束' + S3 + '运动中的若干问题的指示.docx',
    'footnotes': [
        (S3 + '运动', S3 + '运动，指1951年底至1952年10月在中国党政机关、国营企业等单位开展的' + LQ + '反贪污、反浪费、反官僚主义' + RQ + '运动。'),
        (LQ + '打虎' + RQ, LQ + '打虎' + RQ + '，' + S3 + '运动中的术语，指查处和打击贪污数额巨大的贪污分子（即' + LQ + '大老虎' + RQ + '）。'),
        (LQ + '五反' + RQ, S5 + '运动，指1952年上半年在资本主义工商业者中开展的' + LQ + '反行贿、反偷税漏税、反盗骗国家财产、反偷工减料、反盗窃国家经济情报' + RQ + '的运动。'),
        ('人民法庭', '指' + S3 + '运动中在机关单位内临时设立的审判机构，负责审理贪污案件，作出行政和刑事处分决定。'),
        ('市节约检查委员会', '节约检查委员会，' + S3 + '运动期间设立的领导机构，中央设中央节约检查委员会（薄一波任主任），各地设地方节约检查委员会。'),
        (LQ + '双皮坐探' + RQ, LQ + '双皮坐探' + RQ + '，指同时具有国民党特务或反革命组织背景、又受资本家指派打入国家经济部门刺探经济情报的人员。' + LQ + '双皮' + RQ + '指其兼具反革命和资产阶级双重身份。'),
        ('惩治反革命条例', '即1951年2月21日中央人民政府公布的《中华人民共和国惩治反革命条例》，共21条，是镇压反革命运动中的基本法律依据。'),
        ('整编', '整编，指精简机关机构、调整人员编制、提高工作效率的组织整顿工作。'),
        ('民主补课', '民主补课，指在企业民主改革和思想改造运动中对尚未进行或尚未完成民主教育的人员补上民主教育的内容。'),
        ('党员标准的八项条件', '即党员八项标准，1951年3月第一次全国组织工作会议通过的《共产党员标准的八项条件》。'),
    ],
    'end_notes': [
        '【需核】1952年5月20日中央关于争取' + S5 + '斗争胜利结束中的几个问题的指示的具体内容需要核实。',
        '【需核】1952年4月5日中央关于可以不进行交代与国家工作人员关系问题的指示内容需要核查。',
    ],
}

# 03: 与傅作义的谈话
ANNOTATIONS['03'] = {
    'filename': '1952年6月20日，与傅作义的谈话.docx',
    'footnotes': [
        ('傅作义', '傅作义（1895\u20141974），字宜生，山西荣河（今临猗）人。原国民党军华北' + LQ + '剿总' + RQ + '总司令。1949年1月率部接受和平改编，使古都北京免于战火完整保存。时任中央人民政府水利部部长、国防委员会副主席。'),
        ('祈年殿', '祈年殿，位于北京天坛内，是天坛的主体建筑。始建于明永乐十八年（1420年），清乾隆十六年（1751年）重建后改为今名。明清两代皇帝在此举行祈谷大典。'),
        ('天坛', '天坛，位于北京，始建于明永乐十八年（1420年），是明清两代皇帝祭天和祈谷的坛庙，占地约273万平方米，1998年被列为世界文化遗产。'),
    ],
    'end_notes': [
        '【需核】1949年初解放军炮击天坛临时飞机场、炸坏祈年殿一角的具体经过和时间需要核实。',
        '【需核】文末' + LQ + '宜生' + RQ + '应改为' + LQ + '傅作义' + RQ + '或其他称谓，人物对话中称谓的标注规范需要统一。',
    ],
}

# 04: 与陈赓等人的谈话
ANNOTATIONS['04'] = {
    'filename': '1952年6月23日，与陈赓等人的谈话.docx',
    'footnotes': [
        ('陈赓', '陈赓（1903\u20141961），湖南湘乡人，1922年加入中国共产党，黄埔军校第一期毕业。时任中国人民志愿军副司令员兼第三兵团司令员（1952年6月奉调回国组建军事工程学院）。1955年被授予大将军衔。'),
        ('最可爱的人', LQ + '最可爱的人' + RQ + '，对中国人民志愿军的尊称。出自作家魏巍1951年4月11日在《人民日报》发表的报告文学《谁是最可爱的人》。此后' + LQ + '最可爱的人' + RQ + '成为志愿军和人民解放军的光荣代称。'),
        ('军事工程学院', '即中国人民解放军军事工程学院，简称' + LQ + '哈军工' + RQ + '，1953年9月1日在黑龙江省哈尔滨市正式成立。陈赓任首任院长兼政治委员。学院设有空军工程、炮兵工程、海军工程等系，是新中国第一所综合性高等军事工程技术院校。'),
        ('红军学校', '即中国工农红军学校，1931年秋在江西瑞金创办，先后由萧劲光、刘伯承、叶剑英等任校长，培养了大批红军军事政治干部。'),
        ('红军干部团', '红军干部团，1934年10月中央红军长征前夕组建，由红军大学和四个步兵学校合编而成，陈赓任团长，宋任穷任政治委员。在长征中发挥了重要作用。'),
        ('坑道', '坑道作战，抗美援朝战争中中国人民志愿军创造的重要战术。1952年春，志愿军在朝鲜战场上大规模构筑坑道工事，形成以坑道为骨干的防御体系，有效应对了美军的空中优势和炮火打击。'),
    ],
    'end_notes': [
        '【需核】陈赓在1952年6月从前线调回的具体时间，以及' + LQ + '志愿军副司令员' + RQ + '是否为当时全部职务。',
        '【需核】斯大林建议毛泽东建立军事工程学院的具体时间和背景，建议核查相关外交资料。',
    ],
}

# 05: 富农成份党员党籍
ANNOTATIONS['05'] = {
    'filename': '1952年6月9日，关于处理农村中富农成份的党员的党籍问题的新规定.docx',
    'footnotes': [
        ('富农', '富农，中国农村阶级成分之一，指占有较多土地和生产资料，自己参加劳动但主要依靠剥削雇佣劳动为生的农民。在土地改革和农业合作化运动中属于限制和逐步消灭的对象。'),
        ('华东局组织部', '中共中央华东局组织部，负责华东地区党的干部管理和组织建设工作。'),
        ('东北局', '中共中央东北局，中共中央在东北地区的代表机关，1945年成立，辖辽宁、吉林、黑龙江三省党的工作。'),
        ('土地改革', '即土地改革，新中国成立初期在农村进行的废除封建土地所有制的革命运动。至1952年，除部分少数民族地区外，全国土地改革已基本完成。'),
        (LQ + '组织起来' + RQ, LQ + '组织起来' + RQ + '，毛泽东1943年11月29日在陕甘宁边区劳动英雄大会上发表讲话《组织起来》，提出' + LQ + '把群众力量组织起来' + RQ + '是农民由穷变富的必由之路。此后' + LQ + '组织起来' + RQ + '成为中国农业合作化运动的核心口号。'),
        ('农业集体化', '农业集体化，指将个体农民经济改造为社会主义集体经济的运动。中国的农业集体化从互助组开始，经过初级农业生产合作社，到1956年基本完成高级农业生产合作社的建立。'),
        ('高利贷者', '高利贷者，指以超过正常利率的高利息发放贷款牟利的人。新中国成立后，高利贷活动被明令禁止。'),
    ],
    'end_notes': [
        '【需核】1949年7月中央组织部复东北局组织部关于' + LQ + '暂保留其党籍' + RQ + '的规定原文内容需要核实。',
        '【需核】华东局组织部5月10日来函及附件的具体内容需要核查。',
    ],
}

# 06: 体育题词
ANNOTATIONS['06'] = {
    'filename': '1952年6月，为中华全国体育总会第二届代表大会及北京第一届体育运动会题词.docx',
    'footnotes': [
        ('中华全国体育总会', '中华全国体育总会，简称全国体总，1949年10月26日在北京成立。朱德任名誉主席，是中华人民共和国成立后建立的第一个全国性群众体育组织。其任务是推动群众性体育运动的发展和运动技术水平的提高。'),
    ],
    'end_notes': [],
}


# ═══════ MAIN ═══════
def main():
    all_files = [f for f in os.listdir(MONTH_DIR) if f.endswith('.docx') and not f.startswith('__')]
    for key, data in ANNOTATIONS.items():
        fname = data['filename']
        matched = None
        for af in all_files:
            if af.startswith(fname[:15]): matched = af; break
        if matched: data['filename'] = matched; process_file(data)
        else: print('NOT FOUND: ' + fname[:30])
    for f in os.listdir(MONTH_DIR):
        if f.startswith('__'): os.remove(os.path.join(MONTH_DIR, f))
    print('\n6月初注完成!')

if __name__ == '__main__': main()
