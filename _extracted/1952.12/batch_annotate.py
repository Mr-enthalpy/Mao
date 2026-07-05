#!/usr/bin/env python3
"""12月份批量初注脚本"""
import os,shutil
from docx import Document
from lxml import etree
WNS='http://schemas.openxmlformats.org/wordprocessingml/2006/main';W14NS='http://schemas.microsoft.com/office/word/2010/wordml'
MD=r'E:\1952年大传\1952年大传\1952.12'; OP='_LLM初注'
LQ='\u201c';RQ='\u201d';ANG='\u300a';ANR='\u300b'

FT=('<w:footnote xmlns:w="%s" xmlns:w14="%s" w:id="{id}"><w:p><w:pPr><w:rPr><w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/><w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/><w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/></w:rPr></w:pPr><w:r><w:rPr><w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/><w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/><w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/></w:rPr><w:footnoteRef/></w:r><w:r><w:rPr><w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/><w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/><w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r></w:p></w:footnote>')%(WNS,W14NS)

def esc(s):return s.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')
def gmi(fr):return max(int(n.get('{%s}id'%WNS,'0'))for n in fr.findall('{%s}footnote'%WNS))
def ins(p,s,fid):
    ft=p.text
    if s not in ft:return False
    ep=ft.index(s)+len(s);c=0
    for r in p.runs:
        re=c+len(r.text)
        if c<=ep<=re:
            sp=ep-c;af=r.text[sp:];r.text=r.text[:sp]
            e=etree.fromstring('<w:r xmlns:w="%s"><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr><w:footnoteReference w:id="%d"/></w:r>'%(WNS,fid))
            r._element.addnext(e)
            if af:e.addnext(etree.fromstring('<w:r xmlns:w="%s"><w:rPr></w:rPr><w:t xml:space="preserve">%s</w:t></w:r>'%(WNS,af)))
            return True
        c=re
    return False
def afn(fp,fid,tx):
    r=etree.fromstring(fp.blob);r.append(etree.fromstring(FT.format(id=fid,text=esc(tx))))
    fp._blob=etree.tostring(r,xml_declaration=True,encoding='UTF-8',standalone=True)
def aend(d,it):
    if not it:return
    d.add_paragraph().add_run('\n'+'\u2500'*40)
    rr=d.add_paragraph().add_run('本篇待人工确认事项：');rr.bold=True
    for i,item in enumerate(it,1):d.add_paragraph().add_run('%d. %s'%(i,item))
def proc(data):
    fn=data['filename'];src=os.path.join(MD,fn)
    dst=os.path.join(MD,fn.replace('.docx','_'+OP+'.docx'))
    if os.path.exists(dst):print('SKIP');return
    shutil.copy2(src,dst);doc=Document(dst)
    fp=None
    for rel in doc.part.rels.values():
        if 'footnote' in rel.reltype.lower():fp=rel.target_part;break
    if fp is None:print('ERR');return
    nid=gmi(etree.fromstring(fp.blob))+1
    print('>>> '+fn[:30]+' (id=%d)'%nid)
    ic=0
    for st,ft in data['fn']:
        for p in doc.paragraphs:
            if st in p.text and ins(p,st,nid):afn(fp,nid,ft);nid+=1;ic+=1;break
        else:print('  NF: '+st[:50])
    if data.get('en'):aend(doc,data['en'])
    doc.save(dst);print('  -> %d'%ic)

A={}
A['01']={'fn':[
('艾森豪威尔','艾森豪威尔（1890\u20141969），美国共和党人。1952年11月当选美国第34任总统，1953年1月就职。此前曾任二战欧洲盟军最高司令、北约武装部队最高司令。'),
('李承晚','李承晚（1875\u20141965），大韩民国首任总统（1948\u20141960在位）。朝鲜战争期间坚持' + LQ + '北进统一' + RQ + '方针，拒绝停战谈判。'),
('通川','通川，位于朝鲜半岛东海岸（今朝鲜江原道通川郡），西濒日本海。1952年冬毛泽东判断此处可能为美军战术性登陆地点之一。'),
('瓮津半岛','瓮津半岛，位于朝鲜西海岸黄海道南端，突出于黄海之中，与三八线南侧的汉江口相对。1952年冬毛泽东判断此处可能为美军战术性登陆地点。'),
('梦金浦里半岛','梦金浦里半岛，位于朝鲜西海岸瓮津半岛以北，突出于黄海。1952年冬毛泽东判断此处可能为美军战术性登陆地点之一。'),
('元山','元山，朝鲜东海岸最大港口城市（今朝鲜江原道首府），濒临永兴湾。朝鲜战争期间为重要战略要地，美军曾多次轰炸该港。'),
('咸兴','咸兴，朝鲜咸镜南道首府，朝鲜第二大城市和重要工业基地，位于东海岸。朝鲜战争期间美军曾于1950年10月在元山、咸兴一带实施登陆。'),
('镇南浦','镇南浦（今南浦），朝鲜西海岸重要港口城市，位于大同江入海口，是拱卫首都平壤的海上门户。'),
('新南州','新南州（亦译新安州），朝鲜平安南道城市，位于清川江畔，是平壤以北的重要交通枢纽。'),
('铁山半岛','铁山半岛，位于朝鲜西北部平安北道，突出于西朝鲜湾。1952年冬毛泽东判断此处为美军可能实施战略性登陆的地点之一。'),
('安东','安东（今丹东），位于辽宁省东南部鸭绿江畔，与朝鲜新义州隔江相望。抗美援朝期间为中国军队入朝和物资运输的主要枢纽城市。'),
('李弥','李弥（1902\u20141973），云南盈江人，国民党陆军中将。1950年率残部逃入缅甸北部，在台湾当局支持下继续从事反共军事活动，多次骚扰云南边境。'),
('金化','金化，位于朝鲜江原道中部（今属朝鲜）。1952年10月至11月，中国人民志愿军在此以北的上甘岭地区与美军进行了著名的上甘岭战役。'),
('卡秋莎','卡秋莎' + RQ + LQ + '，即BM-13型多管火箭炮，苏联在二战中研制并大量使用的自行火箭炮，以发射时特有的呼啸声闻名。中国人民志愿军在朝鲜战场广泛装备使用。'),
('莫洛托夫','莫洛托夫（1890\u20141986），时任苏联部长会议第一副主席兼外交部长，斯大林最亲密的助手之一。'),
('叶季壮','叶季壮（1893\u20141967），广东新兴人，时任中央人民政府对外贸易部部长，1952年赴莫斯科负责对苏贸易谈判。'),],
'filename':'1952年12月17日，关于朝鲜战争形势等问题致斯大林电——小楠同志.docx','en':[
    '【需核】' + LQ + '通川' + RQ + LQ + '梦金浦里半岛' + RQ + LQ + '新南州' + RQ + LQ + '铁山半岛' + RQ + '等地名在朝鲜战争文献中的标准译法。',
    '【需核】电报中所列火炮及炮弹型号的具体数字和口径规格。',
    '【需核】' + LQ + '汉川' + RQ + '具体位置待核——该地名在朝鲜西海岸的具体坐标和今名。',
]}

A['02']={'fn':[
('边打、边稳、边建','' + LQ + '边打、边稳、边建' + RQ + '是中共中央在朝鲜战争期间提出的经济建设方针：一边进行抗美援朝战争，一边稳定国内物价和市场，一边开展国家经济建设。1952年8月由陈云首次概括提出。'),
('鞍钢','鞍钢，即鞍山钢铁公司，位于辽宁省鞍山市，始建于1916年，是当时中国最大的钢铁工业基地。1952年时被列为' + LQ + '一五' + RQ + '计划重点建设项目的核心工程。'),],
'filename':'1952年12月22日，关于编制一九五三年计划及五年建设计划纲要的指示.docx','en':[]}

A['03']={'fn':[
('叶子龙','叶子龙（1916\u20142003），湖南浏阳人。1935年起长期担任毛泽东机要秘书，负责毛泽东的日常文件处理和通信联络，直至1962年。'),],
'filename':'1952年12月24日，与叶子龙的谈话.docx','en':[]}

A['04']={'fn':[
('王鹤滨','王鹤滨（1924\u20142018），河北安国人。1949年至1953年担任毛泽东的保健医生兼生活秘书，负责毛泽东的日常饮食起居和医疗保健。'),],
'filename':'1952年12月26日，与王鹤滨的谈话.docx','en':[]}

A['05']={'fn':[
('全总党组','全总，指中华全国总工会，中国共产党领导下的全国工会组织的领导机关，1925年成立。全总党组即中国共产党在全国总工会的领导核心组织。'),
('五三工厂','五三工厂，位于沈阳的军工模范工厂，前身为日伪时期的奉天造兵所。新中国成立后以其政治工作与经济工作相结合的管理经验闻名全国，被誉为全国工业战线的先进典型。'),
('中财委','中财委，即中央人民政府政务院财政经济委员会，1949年10月成立，陈云任主任，是新中国成立初期国家经济工作的最高领导机构。1954年撤销。'),],
'filename':'1952年12月27日，关于工会基层工作会议的总结报告.docx','en':[]}

A['06']={'fn':[
('汉川鸭绿江线','汉川鸭绿江线，指朝鲜西海岸从汉川至鸭绿江口的防御地带。毛泽东将其列为防止美军登陆的一号危险区，强调' + LQ + '决不能许敌在西海岸登陆' + RQ + '。'),
('通川元山线','通川元山线，指朝鲜东海岸从通川至元山的防御地带。毛泽东判断此线为美军第二个可能的登陆区域。'),
('镇南浦汉川线','镇南浦汉川线，指朝鲜西海岸从镇南浦（今南浦）至汉川的防御地带。毛泽东判断此线为美军第三个可能的登陆区域。'),],
'filename':'1952年12月4日，在朝鲜战争形势与明年作战方针任务的报告上的重要批示.docx','en':[
    '【需核】' + LQ + '汉川' + RQ + '在朝鲜的具体位置和今名。',
    '【需核】此批示原文出处及完整版本文献来源。',
]}

A['07']={'fn':[
('民族区域自治','民族区域自治，中国共产党处理国内民族问题的基本政策。在少数民族聚居地区建立自治地方，由少数民族当家作主管理本民族内部事务。1952年8月中央人民政府颁布《民族区域自治实施纲要》。'),
('民族民主联合政府','民族民主联合政府，建国初期在民族杂居地区建立的一种过渡性政权形式，由各民族代表联合组成政府。后逐步过渡为民族区域自治地方。'),
('民族学院','民族学院，专门培养少数民族干部和各类专业人才的高等院校。1950年中央民族学院（今中央民族大学）在北京成立，此后西北、西南、中南等地相继建立地方民族学院。'),
('驱梅站','驱梅站，即梅毒防治站。新中国成立初期在少数民族地区开展大规模性病防治工作，' + LQ + '驱梅' + RQ + '（即驱除梅毒）是当时少数民族地区卫生工作的重点之一。'),
('中央民族事务委员会','中央民族事务委员会，简称' + LQ + '中央民委' + RQ + '，1949年10月成立，主管全国民族事务，首任主任委员为李维汉。'),],
'filename':'1952年12月7日，关于制订五年建设计划应重视少数民族地区建设的指示.docx','en':[]}

A['08']={'fn':[
('杨步浩','杨步浩（1905\u20141977），陕西横山人，陕甘宁边区著名劳动英雄。1943年在边区大生产运动中被评为劳动模范，曾代毛泽东耕种，与毛泽东交谊甚厚。'),
('胡宗南','胡宗南（1896\u20141962），浙江孝丰人，国民党陆军一级上将，黄埔军校一期。1947年率部进攻并占领延安，1949年败退四川后逃往台湾。'),],
'filename':'1952年12月，与杨步浩的谈话.docx','en':[]}

A['09']={'fn':[
('高智','高智，陕西佳县人，1952年至1962年担任毛泽东的机要秘书，负责处理毛泽东的日常文件和通信。'),
('佳县','佳县，位于陕西省东北部，黄河西岸，属榆林市管辖。1947年毛泽东转战陕北时曾途经佳县，题词' + LQ + '站在最大多数劳动人民的一面' + RQ + '。'),
('绥德师范','绥德师范，全称陕西省绥德师范学校，位于陕北绥德县。1923年由李子洲创办，是陕北地区历史最悠久的中等师范学校之一，为陕北革命培养了大量干部。'),
('霍去病','霍去病（前140\u2014前117），西汉名将，河东平阳（今山西临汾）人。官至骠骑将军、大司马，封冠军侯。六次出击匈奴，战功卓著，24岁病逝。'),
('霍碧英','霍碧英，高智的妻子，当时在中办机要局工作。'),
('中央机要局','中央机要局，即中共中央办公厅机要局，负责中共中央和中央领导人的密码通信及机要文电处理工作。'),],
'filename':'1952年12月，与高智的谈话.docx','en':[
    '【需核】高智生卒年及详细履历。',
    '【需核】霍碧英的生平信息。',
]}

A['10']={'fn':[
('延安地委','延安地委，即中国共产党延安地方委员会，1950年成立，管辖延安地区各县党务工作。延安是1937年至1947年中共中央所在地。'),
('杨劳动英雄','杨劳动英雄，指杨步浩，陕甘宁边区劳动模范。详见同日前后' + LQ + '与杨步浩的谈话' + RQ + '篇脚注。'),
('六合乡','六合乡，位于延安地区，杨步浩的家乡。'),],
'filename':'1952年12月，给延安地委的信（部分）.docx','en':[]}

def main():
    allf=[f for f in os.listdir(MD)if f.endswith('.docx')and not f.startswith('__')]
    for k,d in A.items():
        fn=d['filename'];m=None
        for af in allf:
            if af.startswith(fn[:20]):m=af;break
        if m:d['filename']=m;proc(d)
        else:print('NF: '+fn[:30])
    for f in os.listdir(MD):
        if f.startswith('__'):os.remove(os.path.join(MD,f))
    print('\n12月初注完成!')

if __name__=='__main__':main()
