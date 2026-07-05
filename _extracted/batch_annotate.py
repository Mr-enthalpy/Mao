#!/usr/bin/env python3
"""
批量初注脚本：为 1952 年 1 月份文档插入脚注和 Word 批注。
输出：原文件名_LLM初注.docx
"""
import os, shutil, re, copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W14NS = 'http://schemas.microsoft.com/office/word/2010/wordml'
RNS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

JAN_DIR = r'E:\1952年大传\1952年大传\1952.1'
OUT_PREFIX = '_LLM初注'

# ───────────────── footnotes template XML ─────────────────
FOOTNOTE_TPL = (
    '<w:footnote xmlns:w="%s" xmlns:w14="%s" w:id="{id}">'
    '<w:p>'
    '<w:pPr>'
    '<w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr>'
    '</w:pPr>'
    '<w:r>'
    '<w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr>'
    '<w:footnoteRef/>'
    '</w:r>'
    '<w:r>'
    '<w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr>'
    '<w:t xml:space="preserve">{text}</w:t>'
    '</w:r>'
    '</w:p>'
    '</w:footnote>'
) % (WNS, W14NS)


def get_max_footnote_id(fn_root):
    """Get the max footnote ID to determine next available."""
    max_id = 0
    for fn in fn_root.findall('{%s}footnote' % WNS):
        fid = int(fn.get('{%s}id' % WNS, '0'))
        if fid > max_id:
            max_id = fid
    return max_id


def insert_footnote_ref(run_elem, fn_id):
    """Insert a footnote reference run after the given run element."""
    ref_xml = (
        '<w:r xmlns:w="%s">'
        '<w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr>'
        '<w:footnoteReference w:id="%d"/>'
        '</w:r>'
    ) % (WNS, fn_id)
    ref_elem = etree.fromstring(ref_xml)
    run_elem.addnext(ref_elem)
    return ref_elem


def insert_after_text(paragraph, search_text, fn_id):
    """Find search_text in runs and insert footnote ref after it.
    Returns True if inserted."""
    full_text = paragraph.text
    if search_text not in full_text:
        return False

    pos_in_full = full_text.index(search_text)
    end_pos = pos_in_full + len(search_text)

    # Find which run contains the end of the match
    current_pos = 0
    for run in paragraph.runs:
        run_len = len(run.text)
        run_end = current_pos + run_len

        if current_pos <= end_pos <= run_end:
            # Split at the insertion point
            split_at = end_pos - current_pos
            before = run.text[:split_at]
            after = run.text[split_at:]
            run.text = before

            # Insert footnote reference
            ref_elem = insert_footnote_ref(run._element, fn_id)

            # Add trailing text if any
            if after:
                after_xml = (
                    '<w:r xmlns:w="%s">'
                    '<w:rPr></w:rPr>'
                    '<w:t xml:space="preserve">%s</w:t>'
                    '</w:r>'
                ) % (WNS, after)
                after_elem = etree.fromstring(after_xml)
                ref_elem.addnext(after_elem)
            return True

        current_pos = run_end

    return False


def add_footnote_to_part(fn_part, fn_id, fn_text):
    """Add a footnote to the footnotes part XML."""
    fn_root = etree.fromstring(fn_part.blob)
    fn_xml = FOOTNOTE_TPL.format(id=fn_id, text=fn_text)
    fn_elem = etree.fromstring(fn_xml)
    fn_root.append(fn_elem)
    fn_part._blob = etree.tostring(fn_root, xml_declaration=True,
                                     encoding='UTF-8', standalone=True)


def get_comments_part(doc):
    """Get or create the comments part."""
    for rel in doc.part.rels.values():
        if 'comment' in rel.reltype.lower():
            return rel.target_part, rel
    return None, None


def add_comment(doc, doc_body, search_text, comment_text):
    """Add a Word comment at the position of search_text."""
    # Find the text in document body (XML level)
    full_text_parts = []
    for p in doc_body.iter('{%s}p' % WNS):
        t_texts = []
        for t in p.iter('{%s}t' % WNS):
            if t.text:
                t_texts.append(t.text)
        full_text_parts.append(''.join(t_texts))

    # Not implementing full comment insertion for now - too complex.
    # Instead, add as a special footnote-like comment at end of doc.
    # For the initial annotation, we mark uncertain items in the end-of-doc section.
    pass


def add_end_note_paragraph(doc, items):
    """Add '本篇待人工确认事项' paragraph at end of document."""
    if not items:
        return
    para = doc.add_paragraph()
    run = para.add_run('\n\n本篇待人工确认事项：')
    run.bold = True
    for i, item in enumerate(items, 1):
        para = doc.add_paragraph()
        run = para.add_run('%d. %s' % (i, item))


def add_comment_paragraph(doc, search_text, comment_label, comment_content):
    """At the end of the current paragraph, add a comment marker paragraph."""
    # Simplified: add at end of document as italic note
    # The proper way would be insertion at the text position, but
    # Word comments require complex XML manipulation.
    # For now, we append batch comments at end.

    # We'll collect comments and add them to the end-of-doc section
    pass


# ───────────────── ANNOTATION DATA ─────────────────
# Each file: (filename, [annotations])
# Annotation: (search_text, footnote_text, type)
# type: 'footnote' or 'comment'
# For comments: comment_label (e.g., '【需核】')

ANNOTATIONS = {}

# File 01: 与华罗庚的谈话
ANNOTATIONS['01'] = {
    'filename': '1952年1月1日，与华罗庚的谈话.docx',
    'footnotes': [
        ('梁思成', '梁思成（1901—1972），广东新会人，中国近代建筑学家、建筑教育家。时任清华大学建筑系主任，参与北京城市规划工作。'),
    ],
    'comments': [],
    'end_notes': [],
}

# File 02: 与罗光禄的谈话
ANNOTATIONS['02'] = {
    'filename': '1952年1月1日，与罗光禄的谈话.docx',
    'footnotes': [
        ('罗光禄', '罗光禄（1919\u20141994），四川苍溪人，1933年参加中国工农红军。长期担任毛泽东的机要秘书，负责文件收发、归档等机要工作。'),
        ('尚昆', '即杨尚昆（1907—1998），四川潼南人，1926年加入中国共产党。时任中共中央办公厅主任。'),
        ('谭政', '谭政（1906—1988），湖南湘乡人，1927年参加湘赣边界秋收起义。时任中南军区第三政治委员。'),
    ],
    'comments': [],
    'end_notes': [],
}

# File 03: 宣传文教部门三反指示
# Note: The document uses Chinese quotation marks \u201c (") and \u201d (")
LQ = '\u201c'  # "
RQ = '\u201d'  # "
ANNOTATIONS['03'] = {
    'filename': '1952年1月22日，关于宣传文教部门应无例外地进行\u201c三反\u201d运动的指示.docx',
    'footnotes': [
        (LQ + '三反' + RQ + '运动', LQ + '三反' + RQ + '运动，指1951年底至1952年10月在中国党政机关、国营企业等单位开展的' + LQ + '反贪污、反浪费、反官僚主义' + RQ + '运动。毛泽东于1951年12月发出号召，运动随即在全国展开。'),
        ('西北局', '中共中央西北局，中共中央在西北地区的代表机关。1941年5月成立，驻延安（后迁西安），辖陕西、甘肃、宁夏、青海、新疆五省区党的工作，1954年12月撤销。'),
        (LQ + '大老虎' + RQ, LQ + '大老虎' + RQ + '，' + LQ + '三反' + RQ + '运动中的特有名词，指贪污数额巨大（旧人民币1亿元以上，约合1955年币制改革后的新人民币1万元）的贪污分子，是运动的重点打击对象。'),
        ('坦白检举', '坦白检举，' + LQ + '三反' + RQ + '运动中发动群众的基本方式。' + LQ + '坦白' + RQ + '指有问题的干部主动交代自己的贪污、浪费或官僚主义问题；' + LQ + '检举' + RQ + '指群众揭发他人的问题。'),
        ('中苏友协分会', '中苏友好协会，1949年10月5日成立于北京，是以促进中苏两国人民友谊和文化交流为宗旨的全国性群众团体。中央设总会，各地设分会。'),
        ('抗美援朝分会', '中国人民抗美援朝总会的地方分会。总会于1950年10月26日在北京成立，郭沫若任主席，负责动员和组织全国人民支援抗美援朝战争。'),
        ('1亿元', '此处' + LQ + '1亿元' + RQ + '指旧人民币。1955年3月1日中国人民银行发行新人民币，以1:10000的比例收兑旧人民币。1亿元旧币合新人民币1万元。'),
        ('党刊', '指中国共产党各级组织主办的党内刊物，用于传达中央指示、交流工作经验和指导实际工作。'),
    ],
    'comments': [],
    'end_notes': [
        '【需核】文中"西北局1月11日所发西北宣传文化教育部门三反运动情况的报告"的具体保存情况和建议出处需要核实。',
    ],
}

# File 04: 给斯大林的电报
ANNOTATIONS['04'] = {
    'filename': '1952年1月31日，给斯大林的电报.docx',
    'footnotes': [
        ('朝鲜停战谈判', '朝鲜停战谈判，指1951年7月10日开始在开城（后移至板门店）举行的朝鲜战争交战双方为结束战争而进行的谈判。至1952年1月，双方已就军事分界线、停火停战安排等大部分议程达成协议，但在战俘遣返问题上陷入僵局。'),
        ('军事分界线', '军事分界线，朝鲜停战谈判中双方同意在朝鲜半岛划定的、以双方实际接触线为基础的军事界线。谈判中双方已就军事分界线的划定达成原则协议。'),
        ('非军事区', '非军事区，指根据停战协议在军事分界线两侧各后退一定距离建立的缓冲地带，禁止任何军事行动和武装人员进入。最终协议将此距离定为2公里。'),
        ('战俘处理', '战俘处理，朝鲜停战谈判第四项议程，即战俘遣返问题。这是谈判中争论最激烈的议题，中方主张按照《关于战俘待遇之日内瓦公约》全部遣返，美方则主张"自愿遣返"原则。'),
        ('自愿遣返', '"自愿遣返"，朝鲜停战谈判中美方提出的战俘遣返原则，主张由战俘本人自由选择是否遣返。中方认为这违反了1949年《关于战俘待遇之日内瓦公约》中"战争结束后应迅速释放与遣返"的规定。'),
        ('中立国', '中立国观察机构，朝鲜停战谈判中提议设立的、由未参加朝鲜战争的中立国代表组成的监督机构，负责监督停战协议的执行情况，特别是监督双方不得从外部向朝鲜增派军事人员和装备。'),
        ('停战委员会', '停战委员会，即军事停战委员会，由朝鲜战争交战双方代表组成的联合机构，负责监督停战协议的执行、处理违约指控等。'),
        ('汉城', '汉城，韩国首都，今称首尔（2005年更名）。朝鲜战争期间曾四次易手，此时由韩国控制。'),
        ('仁川', '仁川，韩国西北部港口城市，靠近汉城，1950年9月联合国军在此登陆。'),
        ('釜山', '釜山，韩国东南部重要港口城市，朝鲜战争初期为韩国临时首都所在地。'),
        ('平壤', '平壤，朝鲜民主主义人民共和国首都。朝鲜战争期间曾一度被联合国军占领，后收复。'),
        ('元山', '元山，朝鲜东部重要港口城市，位于朝鲜半岛东海岸中部。'),
        ('清津', '清津，朝鲜东北部工业城市和重要港口，位于咸镜北道。'),
        ('新义州', '新义州，朝鲜西北部边境城市，与中国安东（今丹东）隔鸭绿江相望，是中朝交通要道。'),
    ],
    'comments': [],
    'end_notes': [
        '【需核】朝鲜南北两方共27处地名，目前仅对主要城市加注，其余如襄阳、原州、大邱、江界、咸兴等是否全部加注，请审稿者决定。',
        '【需核】文中"葫口岛"是否为朝鲜实际地名，未能查证到可靠来源。',
    ],
}

# File 05: 立即抓紧三反斗争
ANNOTATIONS['05'] = {
    'filename': '1952年1月4日，关于立即抓紧\u201c三反\u201d斗争的指示.docx',
    'footnotes': [
        (LQ + '三反' + RQ + '斗争', LQ + '三反' + RQ + '斗争，即' + LQ + '三反' + RQ + '运动，指1951年底至1952年10月开展的' + LQ + '反贪污、反浪费、反官僚主义' + RQ + '运动。'),
        ('中央人民政府各党组', '指中央人民政府各部、委、院、署等机关中设立的中国共产党党组，负责在政府机构中贯彻党的方针政策。'),
        ('志愿军', '中国人民志愿军，1950年10月组成并开赴朝鲜参加抗美援朝战争。彭德怀任司令员兼政治委员。'),
        ('薄一波', '薄一波（1908—2007），山西定襄人，1925年加入中国共产党。时任中央人民政府财政部部长、中央节约检查委员会主任，受中央委托主抓' + LQ + '三反' + RQ + '运动。'),
        ('安子文', '安子文（1909—1980），陕西子洲人，1927年加入中国共产党。时任中共中央组织部副部长、中央节约检查委员会委员。'),
        ('中央直属总党委', '即中共中央直属机关总党委，负责中共中央直属机关党的工作。'),
        ('中央节约检查委员会', '中央节约检查委员会，1951年12月成立，薄一波任主任，负责领导全国的增产节约和' + LQ + '三反' + RQ + LQ + '五反' + RQ + '运动。'),
        ('团拜会', '团拜会，指机关团体在新年时举行的集体拜年仪式。此处指1952年元旦中央机关举办的团拜活动。'),
        ('党刊', '指中国共产党各级组织主办的党内刊物。'),
        ('坦白检举', '坦白检举，' + LQ + '三反' + RQ + '运动中发动群众的基本方式。' + LQ + '坦白' + RQ + '指有问题的干部主动交代问题；' + LQ + '检举' + RQ + '指群众揭发他人的贪污、浪费或官僚主义问题。'),
    ],
    'comments': [],
    'end_notes': [
        '【需核】文中提到的薄一波给毛主席的信和中央财政部党组第四号报告的具体内容，建议核实。',
    ],
}

# File 06: 给班禅的电报
ANNOTATIONS['06'] = {
    'filename': '1952年1月下旬，给班禅的电报.docx',
    'footnotes': [
        ('班禅', '即班禅额尔德尼·确吉坚赞（1938—1989），西藏喇嘛教格鲁派（黄教）两大活佛之一。1949年10月1日致电毛泽东表示拥护中央人民政府。1951年底自青海启程返藏，于1952年4月抵达拉萨。'),
        ('达赖喇嘛', '即达赖喇嘛·丹增嘉措（1935—），西藏喇嘛教格鲁派（黄教）两大活佛之一。1951年西藏和平解放后，任西藏自治区筹备委员会主任（1956年）。'),
        ('和平解放西藏办法的协议', '即《中央人民政府和西藏地方政府关于和平解放西藏办法的协议》，又称"十七条协议"，1951年5月23日在北京签订，宣告西藏和平解放。'),
        ('西藏', '西藏，中国西南部的一个自治区。1951年5月23日签订"十七条协议"后和平解放。'),
    ],
    'comments': [],
    'end_notes': [
        '【需核】本电文的具体发送日期（1月下旬哪一天）需要核实。',
    ],
}


# ───────────────── PROCESSING ─────────────────

def process_file(annotation_data):
    """Process a single file: find, insert footnotes, save."""
    fname = annotation_data['filename']
    src = os.path.join(JAN_DIR, fname)
    out_name = fname.replace('.docx', '_' + OUT_PREFIX + '.docx')
    dst = os.path.join(JAN_DIR, out_name)

    # Skip if output already exists
    if os.path.exists(dst):
        print('SKIP (exists): ' + out_name)
        return

    shutil.copy2(src, dst)
    doc = Document(dst)

    # Get footnotes part
    fn_part = None
    for rel in doc.part.rels.values():
        if 'footnote' in rel.reltype.lower():
            fn_part = rel.target_part
            break

    if fn_part is None:
        print('ERROR: No footnotes part in ' + fname)
        return

    fn_root = etree.fromstring(fn_part.blob)
    next_id = get_max_footnote_id(fn_root) + 1
    print('Processing: ' + fname + ' (next_id=' + str(next_id) + ')')

    inserted = 0
    skipped = 0

    for search_text, fn_text in annotation_data['footnotes']:
        fn_text_escaped = fn_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')
        found = False

        for para in doc.paragraphs:
            if search_text in para.text:
                # Verify this is in body text, not in an existing footnote area
                # (paragraphs don't include footnotes in python-docx)
                if insert_after_text(para, search_text, next_id):
                    add_footnote_to_part(fn_part, next_id, fn_text_escaped)
                    # Re-parse fn_root for next insertions
                    fn_root = etree.fromstring(fn_part.blob)
                    next_id += 1
                    max_id = get_max_footnote_id(etree.fromstring(fn_part.blob))
                    inserted += 1
                    found = True
                    break

        if not found:
            skipped += 1
            print('  NOT FOUND: ' + search_text)

    # Add end notes
    if annotation_data.get('end_notes'):
        end_items = annotation_data['end_notes']
        # Add a separator paragraph
        sep = doc.add_paragraph()
        sep_run = sep.add_run('─' * 40)
        sep_run.font.size = doc.styles['Normal'].font.size

        title_para = doc.add_paragraph()
        title_run = title_para.add_run('本篇待人工确认事项：')
        title_run.bold = True

        for i, item in enumerate(end_items, 1):
            p = doc.add_paragraph()
            p_run = p.add_run('%d. %s' % (i, item))

    doc.save(dst)
    print('  -> ' + str(inserted) + ' footnotes inserted, ' + str(skipped) + ' skipped')
    return dst


def main():
    all_files = sorted([f for f in os.listdir(JAN_DIR) if f.endswith('.docx') and not f.startswith('__')])

    for key, data in ANNOTATIONS.items():
        # Match by the first 15+ chars of the filename prefix,
        # ignoring the exact quotes which may differ (curly vs straight)
        fname = data['filename']
        matched = None
        for af in all_files:
            if af.startswith(fname[:15]):
                matched = af
                break
        if matched:
            data['filename'] = matched
            process_file(data)
        else:
            print('FILE NOT FOUND: ' + fname)

    # Clean up test files
    for f in os.listdir(JAN_DIR):
        if f.startswith('__'):
            os.remove(os.path.join(JAN_DIR, f))

    print('\nAll done!')


if __name__ == '__main__':
    main()
