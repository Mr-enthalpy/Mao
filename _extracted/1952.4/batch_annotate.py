#!/usr/bin/env python3
"""4月份批量初注脚本"""
import os, shutil
from docx import Document
from lxml import etree

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W14NS = 'http://schemas.microsoft.com/office/word/2010/wordml'
MONTH_DIR = r'E:\1952年大传\1952年大传\1952.4'
OUT_PREFIX = '_LLM初注'

LQ = '\u201c'
RQ = '\u201d'

FOOTNOTE_TPL = (
    '<w:footnote xmlns:w="%s" xmlns:w14="%s" w:id="{id}">'
    '<w:p><w:pPr><w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr></w:pPr>'
    '<w:r><w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr><w:footnoteRef/></w:r>'
    '<w:r><w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr><w:t xml:space="preserve">{text}</w:t></w:r>'
    '</w:p></w:footnote>'
) % (WNS, W14NS)


def esc(s):
    return s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')


def get_max_footnote_id(fn_root):
    max_id = 0
    for fn in fn_root.findall('{%s}footnote' % WNS):
        max_id = max(max_id, int(fn.get('{%s}id' % WNS, '0')))
    return max_id


def insert_footnote_ref(run_elem, fn_id):
    ref_xml = (
        '<w:r xmlns:w="%s">'
        '<w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr>'
        '<w:footnoteReference w:id="%d"/>'
        '</w:r>'
    ) % (WNS, fn_id)
    return etree.fromstring(ref_xml)


def insert_after_text(paragraph, search_text, fn_id):
    full_text = paragraph.text
    if search_text not in full_text:
        return False
    pos_in_full = full_text.index(search_text)
    end_pos = pos_in_full + len(search_text)
    current_pos = 0
    for run in paragraph.runs:
        run_len = len(run.text)
        run_end = current_pos + run_len
        if current_pos <= end_pos <= run_end:
            split_at = end_pos - current_pos
            after = run.text[split_at:]
            run.text = run.text[:split_at]
            ref_elem = insert_footnote_ref(run._element, fn_id)
            run._element.addnext(ref_elem)
            if after:
                after_xml = (
                    '<w:r xmlns:w="%s"><w:rPr></w:rPr>'
                    '<w:t xml:space="preserve">%s</w:t></w:r>'
                ) % (WNS, after)
                after_elem = etree.fromstring(after_xml)
                ref_elem.addnext(after_elem)
            return True
        current_pos = run_end
    return False


def add_footnote_to_part(fn_part, fn_id, fn_text):
    fn_root = etree.fromstring(fn_part.blob)
    fn_xml = FOOTNOTE_TPL.format(id=fn_id, text=esc(fn_text))
    fn_root.append(etree.fromstring(fn_xml))
    fn_part._blob = etree.tostring(fn_root, xml_declaration=True,
                                     encoding='UTF-8', standalone=True)


def add_end_notes(doc, items):
    if not items:
        return
    sep = doc.add_paragraph()
    sep.add_run('\n' + '\u2500' * 40)
    title_para = doc.add_paragraph()
    run = title_para.add_run('本篇待人工确认事项：')
    run.bold = True
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.add_run('%d. %s' % (i, item))


def process_file(data):
    fname = data['filename']
    src = os.path.join(MONTH_DIR, fname)
    out_name = fname.replace('.docx', '_' + OUT_PREFIX + '.docx')
    dst = os.path.join(MONTH_DIR, out_name)
    if os.path.exists(dst):
        print('SKIP: ' + out_name)
        return
    shutil.copy2(src, dst)
    doc = Document(dst)

    fn_part = None
    for rel in doc.part.rels.values():
        if 'footnote' in rel.reltype.lower():
            fn_part = rel.target_part
            break
    if fn_part is None:
        print('ERROR: no footnotes')
        return

    fn_root = etree.fromstring(fn_part.blob)
    next_id = get_max_footnote_id(fn_root) + 1
    print('>>> ' + fname[:35] + ' (id=%d)' % next_id)

    inserted = 0
    for search_text, fn_text in data['footnotes']:
        found = False
        for para in doc.paragraphs:
            if search_text in para.text:
                if insert_after_text(para, search_text, next_id):
                    add_footnote_to_part(fn_part, next_id, fn_text)
                    fn_root = etree.fromstring(fn_part.blob)
                    next_id += 1
                    inserted += 1
                    found = True
                    break
        if not found:
            print('  NOT FOUND: ' + search_text[:50])

    if data.get('end_notes'):
        add_end_notes(doc, data['end_notes'])

    doc.save(dst)
    print('  -> %d footnotes' % inserted)


# ═════════════════ ANNOTATIONS ═════════════════
ANNOTATIONS = {}

# --- File 01: 关于《矛盾论》的讨论 ---
ANNOTATIONS['01'] = {
    'filename': '1952年4月1日，关于《矛盾论》的讨论志.docx',
    'footnotes': [
        ('\u300a矛盾论\u300b', '《矛盾论》，毛泽东的哲学著作，1937年8月在延安抗日军事政治大学作讲演，后收入《毛泽东选集》第一卷。与《实践论》并称为毛泽东哲学思想的代表作，系统论述了唯物辩证法的对立统一规律。本文中提到的' + LQ + '重新发表' + RQ + '指1952年4月1日《人民日报》重新刊登《矛盾论》。'),
        ('王鹤滨', '王鹤滨（1924\u2014），毛泽东的保健医生，1949年至1954年任毛泽东秘书兼保健医生。'),
        ('菊香书屋', '菊香书屋，位于北京中南海丰泽园内，是毛泽东1949年至1966年的住所和办公室。'),
        ('江青', '江青（1914\u20141991），山东诸城人，毛泽东的夫人。时任中共中央宣传部文艺处副处长。'),
        ('抗美援朝', '抗美援朝，即中国人民志愿军赴朝鲜支援朝鲜人民抗击以美国为首的联合国军的战争，1950年10月开始。'),
        ('美帝国主义', '美帝国主义，指第二次世界大战后以美国为首的帝国主义势力。1949年新中国成立后，美国对中国实行孤立封锁政策，中美两国长期处于对峙状态。'),
    ],
    'end_notes': [
        '【需核】' + LQ + '统一性或同一性' + RQ + '是《矛盾论》中的核心哲学概念（即矛盾的同一性），是否需要另注哲学含义，请人工判断。',
    ],
}

# --- File 02: 给毛泽荣的信 ---
ANNOTATIONS['02'] = {
    'filename': '1952年4月20日，.docx',
    'footnotes': [
        ('泽荣', '毛泽荣（1897\u20141986），毛泽东的堂弟，号冬初，韶山冲人。毛泽东曾多次致信毛泽荣并资助其家庭。'),
    ],
    'end_notes': [
        '【需核】文件名只有日期无标题，是否需要为本文添加标题（如' + LQ + '给毛泽荣的信' + RQ + '）。',
        '【需核】文中毛泽东劝阻毛泽荣' + LQ + '今年不要来京' + RQ + '的具体背景需要核查。',
    ],
}

# --- File 03: 致斯大林电（航空兵）---
ANNOTATIONS['03'] = {
    'filename': '1952年4月22日，.docx',
    'footnotes': [
        ('航空兵', '航空兵，以军用飞机为主要装备的空军兵种，包括歼击航空兵、轰炸航空兵、强击航空兵、运输航空兵等。'),
        ('航校', '即航空学校，培养空军飞行人员和技术人员的中等军事院校。'),
        ('米格-9', '米格-9（МиГ-9），苏联第一代喷气式战斗机，1946年首飞。1950年代初中国曾引进一批米格-9飞机用于训练和作战。'),
        ('米格-15', '米格-15（МиГ-15），苏联后掠翼喷气式战斗机，1947年首飞。抗美援朝战争期间中国人民志愿军空军的主力战机，在与美国F-86的战斗中发挥重要作用。'),
        ('安东', '安东，今辽宁省丹东市，位于鸭绿江畔、中朝边境，是抗美援朝期间中国人民志愿军空军的主要后方基地和物资中转站。1965年更名为丹东。'),
        ('克拉索夫斯基', '克拉索夫斯基（Степан Акимович Красовский），苏联空军元帅，时任苏联驻华军事总顾问，负责协调苏联对中国的军事援助事宜。'),
        ('轰炸机团', '航空兵团是中国人民空军的基本战术单位，一般装备30-40架飞机。轰炸航空兵团主要装备轻型或中型轰炸机。'),
    ],
    'end_notes': [
        '【需核】文件名只有日期无标题，正文标题为' + LQ + '致斯大林电' + RQ + '，是否需要调整为完整标题。',
        '【需核】' + LQ + '歼击机团' + RQ + '与' + LQ + '轰炸机团' + RQ + '的编制飞机数量（' + LQ + '每团31架' + RQ + '）在不同时期可能有所变动，建议核查军史资料。',
    ],
}

# --- File 04: 富农政策复示 ---
ANNOTATIONS['04'] = {
    'filename': '1952年4月29日，.docx',
    'footnotes': [
        ('富农', '富农，中国农村阶级成分之一，指占有较多土地和生产资料，自己参加劳动但主要依靠剥削雇佣劳动（雇工）为生的农民。在土地改革和农业合作化运动中，富农经济属于被限制和逐步消灭的对象。'),
        ('互助组', '互助组，全称农业生产互助组，中国农业合作化的初级形式。由若干农户在个体经济基础上自愿联合，实行集体劳动和换工互助，土地和生产资料仍归各户所有。'),
        ('农业生产合作社', '农业生产合作社，中国农业合作化的中级形式（初级社）。农民将土地入股统一经营，实行按劳分配和按股分红相结合。1953年后快速发展。'),
        ('陕西省委', '即中国共产党陕西省委员会，1949年5月西安解放后由陕甘宁边区的中共组织改组而成。'),
        ('土改', '即土地改革，新中国成立初期在农村进行的废除封建土地所有制的革命运动。至1952年，除部分少数民族地区外，全国土地改革已基本完成。'),
        ('党刊', '指中国共产党各级组织主办的党内刊物，用于传达中央指示和交流工作经验。'),
    ],
    'end_notes': [
        '【需核】文件名只有日期无标题，正文标题为' + LQ + '关于农业互助合作运动中对待富农的政策问题给华东局的复示' + RQ + '，是否需要调整为完整标题。',
        '【需核】华东局4月17日来电的具体内容和陕西省委意见原文需要核查相关档案资料。',
    ],
}

# --- File 05: 与王鹤滨等人的谈话 ---
ANNOTATIONS['05'] = {
    'filename': '1952年4月，与王鹤滨等人的谈话.docx',
    'footnotes': [
        ('王鹤滨', '王鹤滨（1924\u2014），毛泽东的保健医生兼生活秘书，1949年至1954年在中南海工作。'),
        ('菊香书屋', '菊香书屋，位于中南海丰泽园内，毛泽东1949年至1966年的住所。'),
        ('平山县', '平山县，位于河北省西部、太行山东麓，属石家庄市管辖。1948年5月至1949年3月，中共中央驻平山县西柏坡，在此指挥了三大战役并召开了中共七届二中全会。'),
        ('上海医学院', '上海医学院，今复旦大学上海医学院，前身为1927年创办的国立第四中山大学医学院、1932年定名为国立上海医学院，是中国近代著名的医学高等学府。'),
    ],
    'end_notes': [
        '【需核】王鹤滨之子王子冀的姓名用字（' + LQ + '冀' + RQ + '）和出生年份是否需要核实。',
    ],
}


# ═════════════════ MAIN ═════════════════
def main():
    all_files = [f for f in os.listdir(MONTH_DIR) if f.endswith('.docx') and not f.startswith('__')]
    for key, data in ANNOTATIONS.items():
        fname = data['filename']
        matched = None
        for af in all_files:
            if af.startswith(fname[:15]):
                matched = af
                break
        if matched:
            data['filename'] = matched
            process_file(data)
        else:
            print('NOT FOUND: ' + fname[:30])
    for f in os.listdir(MONTH_DIR):
        if f.startswith('__'):
            os.remove(os.path.join(MONTH_DIR, f))
    print('\n4月初注完成!')


if __name__ == '__main__':
    main()
