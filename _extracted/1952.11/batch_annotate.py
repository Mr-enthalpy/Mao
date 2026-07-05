#!/usr/bin/env python3
"""11月份批量初注脚本"""
import os,shutil
from docx import Document
from lxml import etree
WNS='http://schemas.openxmlformats.org/wordprocessingml/2006/main';W14NS='http://schemas.microsoft.com/office/word/2010/wordml'
MD=r'E:\1952年大传\1952年大传\1952.11'; OP='_LLM初注'
LQ='\u201c';RQ='\u201d';ANG='\u300a';ANR='\u300b'

FT=('<w:footnote xmlns:w="%s" xmlns:w14="%s" w:id="{id}"><w:p><w:pPr><w:rPr><w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/><w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/><w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/></w:rPr></w:pPr><w:r><w:rPr><w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/><w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/><w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/></w:rPr><w:footnoteRef/></w:r><w:r><w:rPr><w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/><w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/><w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r></w:p></w:footnote>')%(WNS,W14NS)

def esc(s):return s.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')
def gmi(fr):return max(int(n.get('{%s}id'%WNS,'0'))for n in fr.findall('{%s}footnote'%WNS))
def ins(p,s,fid):
    ft=p.text
    if s not in ft:return False
    ep=ft.index(s)+len(s);c=0
    for r in p.runs:
        re=c+len(r.text)
        if c<=ep<=re:
            sp=ep-c;af=r.text[sp:];r.text=r.text[:sp]
            e=etree.fromstring('<w:r xmlns:w="%s"><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr><w:footnoteReference w:id="%d"/></w:r>'%(WNS,fid))
            r._element.addnext(e)
            if af:e.addnext(etree.fromstring('<w:r xmlns:w="%s"><w:rPr></w:rPr><w:t xml:space="preserve">%s</w:t></w:r>'%(WNS,af)))
            return True
        c=re
    return False
def afn(fp,fid,tx):
    r=etree.fromstring(fp.blob);r.append(etree.fromstring(FT.format(id=fid,text=esc(tx))))
    fp._blob=etree.tostring(r,xml_declaration=True,encoding='UTF-8',standalone=True)
def aend(d,it):
    if not it:return
    d.add_paragraph().add_run('\n'+'\u2500'*40)
    rr=d.add_paragraph().add_run('本篇待人工确认事项：');rr.bold=True
    for i,item in enumerate(it,1):d.add_paragraph().add_run('%d. %s'%(i,item))
def proc(data):
    fn=data['filename'];src=os.path.join(MD,fn)
    dst=os.path.join(MD,fn.replace('.docx','_'+OP+'.docx'))
    if os.path.exists(dst):print('SKIP');return
    shutil.copy2(src,dst);doc=Document(dst)
    fp=None
    for rel in doc.part.rels.values():
        if 'footnote' in rel.reltype.lower():fp=rel.target_part;break
    if fp is None:print('ERR');return
    nid=gmi(etree.fromstring(fp.blob))+1
    print('>>> '+fn[:30]+' (id=%d)'%nid)
    ic=0
    for st,ft in data['fn']:
        for p in doc.paragraphs:
            if st in p.text and ins(p,st,nid):afn(fp,nid,ft);nid+=1;ic+=1;break
        else:print('  NF: '+st[:50])
    if data.get('en'):aend(doc,data['en'])
    doc.save(dst);print('  -> %d'%ic)

A={}
A['01']={'fn':[
('毛宇居','毛宇居（1881\u20141964），毛泽东的堂兄兼塾师，韶山冲人。毛泽东幼年曾随其读书，称其为' + LQ + '宇居大哥' + RQ + '。'),
('韶山学校','韶山学校，位于湖南韶山冲，其前身为韶山小学。1952年毛泽东应乡政府委托亲笔题写校名' + LQ + '韶山学校' + RQ + '。1959年毛泽东回韶山时曾视察该校。'),
('张家山','张家山，位于韶山冲附近的小山丘，韶山小学（后改为韶山学校）新校址所在地。'),],
'filename':'1952年11月10日，与毛宇居的谈话.docx','en':[]}

A['02']={'fn':[
('毛宇居','毛宇居（1881\u20141964），毛泽东的堂兄兼塾师。'),
('谭熙春','谭熙春，韶山冲人，毛泽东的乡亲。具体生平需要进一步查证。'),
('邹普勋','邹普勋，韶山冲人，毛泽东的幼年同学和邻居。其胞姐邹氏是毛泽东的堂婶。'),
('毛碧珠','毛碧珠，又名毛笔珠，韶山冲人，毛泽东的族人。'),
('百万元','此处' + LQ + '百万元' + RQ + '指旧人民币。1955年3月1日中国人民银行发行新人民币，以1:10000的比例收兑旧人民币。文中' + LQ + '两百万元' + RQ + '合新人民币200元，' + LQ + '一百万元' + RQ + '合新人民币100元。'),],
'filename':'1952年11月15日，给毛宇居的便条.docx','en':[
    '【需核】谭熙春、蔚生六婶、毛月秋、邹香庭、张四维等人的准确身份和与毛泽东的关系需要核实。',
]}

A['03']={'fn':[
('汤阴','汤阴，位于河南省北部，古称荡阴，是南宋抗金名将岳飞的故乡。'),
('安阳','安阳，位于河南省最北部，中国八大古都之一。商代后期都城殷墟所在地，也是甲骨文的发现地。'),
('岳飞','岳飞（1103\u20141142），字鹏举，相州汤阴（今河南汤阴）人。南宋杰出的军事家、民族英雄。因坚决主张抗金被秦桧以' + LQ + '莫须有' + RQ + '罪名杀害。其词作《满江红》为千古名篇。'),
('岳庙','岳庙，即岳飞庙，位于汤阴县城内。始建于明景泰元年（1450年），祀奉岳飞。庙内有大量碑碣石刻和岳飞塑像。'),
('满江红','《满江红》，岳飞最著名的词作。上阕' + LQ + '怒发冲冠，凭栏处，潇潇雨歇。抬望眼，仰天长啸，壮怀激烈。' + RQ + '表达了作者抗金救国的英雄气概和壮志未酬的悲愤。'),
('诸葛亮','诸葛亮（181\u2014234），字孔明，琅琊阳都（今山东沂南）人。三国时期蜀汉丞相。'),
('出师表','《出师表》，诸葛亮227年北伐曹魏前写给后主刘禅的奏表，表达了' + LQ + '鞠躬尽瘁，死而后已' + RQ + '的忠心。'),
('殷墟','殷墟，位于安阳市西北郊，是商代后期都城遗址，距今约3300年。1899年在此发现甲骨文，是中国第一个有文献可考、并为考古发掘所证实的古代都城遗址。2006年被列为世界文化遗产。'),
('纣王','纣王（？\u2014前1046年？），子姓，名受，商朝末代君主。在位期间经营东南，统一中原和东夷，但后期荒淫残暴，周武王联合诸侯伐纣，商朝灭亡。'),
('盘庚','盘庚，商朝第二十任君主。在位期间将商都从奄（今山东曲阜）迁至殷（今河南安阳），史称' + LQ + '盘庚迁殷' + RQ + '。此后商朝也被称为殷商。'),
('曹操','曹操（155\u2014220），字孟德，沛国谯（今安徽亳州）人。东汉末年杰出的政治家、军事家、文学家。官渡之战大破袁绍后基本统一北方。曹魏政权的奠基人。'),
('西门豹','西门豹，战国时期魏国人，魏文侯时任邺（今河北临漳西南）令。以破除' + LQ + '河伯娶妇' + RQ + '迷信、开凿十二渠引漳水灌田而著名。'),
('甲骨','甲骨文，刻写在龟甲和兽骨上的文字，是中国已知最早的成体系的文字形式，主要出土于安阳殷墟。1899年被王懿荣首次辨认，至今已发现约15万片。'),
('袁世凯','袁世凯（1859\u20141916），河南项城人，北洋军阀首领。戊戌变法中出卖维新派，后窃取辛亥革命成果任中华民国大总统。1915年12月宣布恢复帝制，次年3月被迫取消，6月病死。'),
('戊戌变法','戊戌变法，1898年（农历戊戌年）以康有为、梁启超为首的维新派通过光绪皇帝推行的资产阶级改良运动，历时103天被慈禧太后发动政变镇压，史称' + LQ + '百日维新' + RQ + '。'),
('袁林','袁林，即袁世凯墓，位于安阳市北郊。1916年至1918年修建，占地9万余平方米，建筑中西合璧。因袁世凯称帝未成，不以' + LQ + '陵' + RQ + '名而以' + LQ + '林' + RQ + '名（谐音）。'),
('徐世昌','徐世昌（1855\u20141939），字卜五，号菊人，天津人。北洋军阀重要人物，1918年至1922年任中华民国大总统。'),],
'filename':'1952年11月1日，在视察汤阴、安阳时的谈话.docx','en':[]}

def main():
    allf=[f for f in os.listdir(MD)if f.endswith('.docx')and not f.startswith('__')]
    for k,d in A.items():
        fn=d['filename'];m=None
        for af in allf:
            if af.startswith(fn[:20]):m=af;break
        if m:d['filename']=m;proc(d)
        else:print('NF: '+fn[:30])
    for f in os.listdir(MD):
        if f.startswith('__'):os.remove(os.path.join(MD,f))
    print('\n11月初注完成!')

if __name__=='__main__':main()
