#!/usr/bin/env python3
"""8月份批量初注脚本"""
import os, shutil
from docx import Document
from lxml import etree

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W14NS = 'http://schemas.microsoft.com/office/word/2010/wordml'
MONTH_DIR = r'E:\1952年大传\1952年大传\1952.8'
OUT_PREFIX = '_LLM初注'
LQ='\u201c'; RQ='\u201d'

FOOTNOTE_TPL = (
    '<w:footnote xmlns:w="%s" xmlns:w14="%s" w:id="{id}">'
    '<w:p><w:pPr><w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr></w:pPr>'
    '<w:r><w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr><w:footnoteRef/></w:r>'
    '<w:r><w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr><w:t xml:space="preserve">{text}</w:t></w:r>'
    '</w:p></w:footnote>'
) % (WNS, W14NS)

def esc(s): return s.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')
def get_max_id(fn_root): return max(int(fn.get('{%s}id'%WNS,'0')) for fn in fn_root.findall('{%s}footnote'%WNS))

def insert_after_text(paragraph, search_text, fn_id):
    full_text = paragraph.text
    if search_text not in full_text: return False
    end_pos = full_text.index(search_text) + len(search_text); cur = 0
    for run in paragraph.runs:
        run_end = cur + len(run.text)
        if cur <= end_pos <= run_end:
            sp = end_pos - cur; after = run.text[sp:]; run.text = run.text[:sp]
            ref = etree.fromstring(('<w:r xmlns:w="%s"><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr><w:footnoteReference w:id="%d"/></w:r>')%(WNS,fn_id))
            run._element.addnext(ref)
            if after:
                a = etree.fromstring(('<w:r xmlns:w="%s"><w:rPr></w:rPr><w:t xml:space="preserve">%s</w:t></w:r>')%(WNS,after))
                ref.addnext(a)
            return True
        cur = run_end
    return False

def add_fn(fn_part, fn_id, fn_text):
    r = etree.fromstring(fn_part.blob)
    r.append(etree.fromstring(FOOTNOTE_TPL.format(id=fn_id, text=esc(fn_text))))
    fn_part._blob = etree.tostring(r, xml_declaration=True, encoding='UTF-8', standalone=True)

def add_end(doc, items):
    if not items: return
    doc.add_paragraph().add_run('\n' + '\u2500' * 40)
    r = doc.add_paragraph().add_run('本篇待人工确认事项：'); r.bold = True
    for i, item in enumerate(items, 1): doc.add_paragraph().add_run('%d. %s' % (i, item))

def process_file(data):
    fname = data['filename']; src = os.path.join(MONTH_DIR, fname)
    dst = os.path.join(MONTH_DIR, fname.replace('.docx','_'+OUT_PREFIX+'.docx'))
    if os.path.exists(dst): print('SKIP'); return
    shutil.copy2(src, dst); doc = Document(dst)
    fn_part = None
    for rel in doc.part.rels.values():
        if 'footnote' in rel.reltype.lower(): fn_part = rel.target_part; break
    if fn_part is None: print('ERROR'); return
    next_id = get_max_id(etree.fromstring(fn_part.blob)) + 1
    print('>>> ' + fname[:30] + ' (id=%d)' % next_id)
    ins = 0
    for st, ft in data['footnotes']:
        for para in doc.paragraphs:
            if st in para.text:
                if insert_after_text(para, st, next_id): add_fn(fn_part, next_id, ft); next_id += 1; ins += 1; break
        else: print('  NF: ' + st[:50])
    if data.get('end_notes'): add_end(doc, data['end_notes'])
    doc.save(dst); print('  -> %d' % ins)

# ═══════ ANNOTATIONS ═══════
ANNOTATIONS = {}

# 01: 司法改革
ANNOTATIONS['01'] = {
    'filename': '1952年8月30日，关于进行司法改革工作应注意的几个问题的指示.docx',
    'footnotes': [
        ('政务院', '政务院，即中央人民政府政务院，1949年10月1日成立，周恩来任总理，是中华人民共和国成立初期的国家政务最高执行机关。1954年9月第一届全国人民代表大会召开后，根据新宪法改称国务院。'),
        ('史良', '史良（1900\u20141985），女，江苏常州人，中国妇女运动领袖、法学家。曾任民盟中央主席。时任中央人民政府司法部部长。'),
        ('旧法观点', LQ + '旧法观点' + RQ + '，指国民党统治时期和西方资本主义国家的法律观念和法律制度，被视为与新中国的人民民主法制相对立的法律观。1952-1953年的司法改革运动以批判和清除旧法观点为核心任务之一。'),
        ('旧司法人员', LQ + '旧司法人员' + RQ + '，指新中国成立前在国民党政府司法机关中任职的审判、检察及其他法律工作人员。1952年司法改革运动中对旧司法人员进行了大规模的组织清理和思想改造。'),
    ],
    'end_notes': [
        '【需核】政务院批准的史良部长关于改造和整顿各级人民法院的报告的具体发布时间和内容需要核实。',
    ],
}

# 02: 全国政协谈话（含程潜）
ANNOTATIONS['02'] = {
    'filename': '1952年8月4日，在全国政协一届常务委员会第三十八次（扩大）会议上的谈话.docx',
    'footnotes': [
        ('全国政协', '中国人民政治协商会议，简称全国政协，1949年9月成立，是中国共产党领导的多党合作和政治协商的重要机构。在1954年全国人民代表大会召开前，代行全国人大的职权。'),
        (LQ + '山字型构造' + RQ, LQ + '山字型构造' + RQ + '，著名地质学家李四光创立的地质力学中的重要概念，指由挤压作用形成的形似' + LQ + '山' + RQ + '字的地质构造体系。毛泽东对地质学有浓厚兴趣，曾多次向李四光请教地质问题。'),
        ('程潜', '程潜（1882\u20141968），字颂云，湖南醴陵人。同盟会会员，国民党元老，曾任国民党第一战区司令长官。1949年8月4日在长沙率部宣布起义，使湖南和平解放。时任中央人民政府委员会委员、湖南省省长。'),
        ('中南海', '中南海，位于北京故宫西侧，中海和南海的合称。1949年后为中共中央和国务院办公所在地，也是毛泽东的住所。'),
        ('南海', '南海，中南海的南海部分。中南海由中海和南海组成，南海位于中海南部，湖中瀛台是乾隆年间修建的皇家园林。'),
        (LQ + '翡翠层楼浮树杪，芙蓉小殿出波心' + RQ, '出自清代文人赞咏中南海的诗句，描写中南海楼阁掩映在绿树之间，亭台仿佛浮出水面的优美景致。' + LQ + '杪' + RQ + '（mi\u01ceo），指树梢。'),
    ],
    'end_notes': [
        '【需核】' + LQ + '大概二三十年吧' + RQ + '这段谈话的上下文和所指内容（毛泽东在回答什么问题）需要核实。',
        '【需核】' + LQ + '山字型构造' + RQ + '这段提问的具体对象（是否李四光在场）和完整语境需要核查。',
    ],
}

# 03: 与梁漱溟的谈话
ANNOTATIONS['03'] = {
    'filename': '1952年8月7日，与梁漱溟的谈话.docx',
    'footnotes': [
        ('梁漱溟', '梁漱溟（1893\u20141988），原名焕鼎，字寿铭，广西桂林人（生于北京）。中国现代哲学家、教育家，乡村建设运动代表人物。1950年后任全国政协委员。1953年因在政协会议上与毛泽东公开争论而受到批判。'),
        ('苏联', '苏维埃社会主义共和国联盟（1922\u20141991），由15个加盟共和国组成的联邦制社会主义国家。新中国成立初期是中苏同盟核心，1950年2月签订《中苏友好同盟互助条约》。'),
    ],
    'end_notes': [
        '【需核】梁漱溟1952年提出去苏联做学术研究的具体背景和时间以及与毛泽东关系的史料需要核查。',
    ],
}

# 04: 给张有成家属唁函
ANNOTATIONS['04'] = {
    'filename': '1952年8月，给张有成家属的唁函.docx',
    'footnotes': [
        ('张有成', '张有成，毛泽东在湖南第一师范学校读书时期的同学和好友。新中国成立后与毛泽东保持通信联系。'),
    ],
    'end_notes': [
        '【需核】张有成的详细生平信息（生卒年、籍贯、与毛泽东交往的具体情况）需要核查。',
    ],
}

def main():
    all_files = [f for f in os.listdir(MONTH_DIR) if f.endswith('.docx') and not f.startswith('__')]
    for key, data in ANNOTATIONS.items():
        fname = data['filename']; matched = None
        for af in all_files:
            if af.startswith(fname[:15]): matched = af; break
        if matched: data['filename'] = matched; process_file(data)
        else: print('NF: ' + fname[:30])
    for f in os.listdir(MONTH_DIR):
        if f.startswith('__'): os.remove(os.path.join(MONTH_DIR, f))
    print('\n8月初注完成!')

if __name__ == '__main__': main()
