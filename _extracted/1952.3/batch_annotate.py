#!/usr/bin/env python3
"""3月份批量初注脚本"""
import os, shutil
from docx import Document
from lxml import etree

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W14NS = 'http://schemas.microsoft.com/office/word/2010/wordml'
MONTH_DIR = r'E:\1952年大传\1952年大传\1952.3'
OUT_PREFIX = '_LLM初注'

LQ = '\u201c'
RQ = '\u201d'

FOOTNOTE_TPL = (
    '<w:footnote xmlns:w="%s" xmlns:w14="%s" w:id="{id}">'
    '<w:p><w:pPr><w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr></w:pPr>'
    '<w:r><w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr><w:footnoteRef/></w:r>'
    '<w:r><w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr><w:t xml:space="preserve">{text}</w:t></w:r>'
    '</w:p></w:footnote>'
) % (WNS, W14NS)


def esc(s):
    return s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')


def get_max_footnote_id(fn_root):
    max_id = 0
    for fn in fn_root.findall('{%s}footnote' % WNS):
        max_id = max(max_id, int(fn.get('{%s}id' % WNS, '0')))
    return max_id


def insert_footnote_ref(run_elem, fn_id):
    ref_xml = (
        '<w:r xmlns:w="%s">'
        '<w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr>'
        '<w:footnoteReference w:id="%d"/>'
        '</w:r>'
    ) % (WNS, fn_id)
    return etree.fromstring(ref_xml)


def insert_after_text(paragraph, search_text, fn_id):
    full_text = paragraph.text
    if search_text not in full_text:
        return False
    pos_in_full = full_text.index(search_text)
    end_pos = pos_in_full + len(search_text)
    current_pos = 0
    for run in paragraph.runs:
        run_len = len(run.text)
        run_end = current_pos + run_len
        if current_pos <= end_pos <= run_end:
            split_at = end_pos - current_pos
            after = run.text[split_at:]
            run.text = run.text[:split_at]
            ref_elem = insert_footnote_ref(run._element, fn_id)
            run._element.addnext(ref_elem)
            if after:
                after_xml = (
                    '<w:r xmlns:w="%s"><w:rPr></w:rPr>'
                    '<w:t xml:space="preserve">%s</w:t></w:r>'
                ) % (WNS, after)
                after_elem = etree.fromstring(after_xml)
                ref_elem.addnext(after_elem)
            return True
        current_pos = run_end
    return False


def add_footnote_to_part(fn_part, fn_id, fn_text):
    fn_root = etree.fromstring(fn_part.blob)
    fn_xml = FOOTNOTE_TPL.format(id=fn_id, text=esc(fn_text))
    fn_root.append(etree.fromstring(fn_xml))
    fn_part._blob = etree.tostring(fn_root, xml_declaration=True,
                                     encoding='UTF-8', standalone=True)


def add_end_notes(doc, items):
    if not items:
        return
    sep = doc.add_paragraph()
    sep.add_run('\n' + '\u2500' * 40)
    title_para = doc.add_paragraph()
    run = title_para.add_run('本篇待人工确认事项：')
    run.bold = True
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.add_run('%d. %s' % (i, item))


def process_file(data):
    fname = data['filename']
    src = os.path.join(MONTH_DIR, fname)
    out_name = fname.replace('.docx', '_' + OUT_PREFIX + '.docx')
    dst = os.path.join(MONTH_DIR, out_name)
    if os.path.exists(dst):
        print('SKIP: ' + out_name)
        return
    shutil.copy2(src, dst)
    doc = Document(dst)

    fn_part = None
    for rel in doc.part.rels.values():
        if 'footnote' in rel.reltype.lower():
            fn_part = rel.target_part
            break
    if fn_part is None:
        print('ERROR: no footnotes')
        return

    fn_root = etree.fromstring(fn_part.blob)
    next_id = get_max_footnote_id(fn_root) + 1
    print('>>> ' + fname[:30] + ' (id=%d)' % next_id)

    inserted = 0
    for search_text, fn_text in data['footnotes']:
        found = False
        for para in doc.paragraphs:
            if search_text in para.text:
                if insert_after_text(para, search_text, next_id):
                    add_footnote_to_part(fn_part, next_id, fn_text)
                    fn_root = etree.fromstring(fn_part.blob)
                    next_id += 1
                    inserted += 1
                    found = True
                    break
        if not found:
            print('  NOT FOUND: ' + search_text[:50])

    if data.get('end_notes'):
        add_end_notes(doc, data['end_notes'])

    doc.save(dst)
    print('  -> %d footnotes' % inserted)


# ═════════════════ ANNOTATIONS ═════════════════
ANNOTATIONS = {}

# --- File 01: 三反运动中党员党内处分规定 ---
ANNOTATIONS['01'] = {
    'filename': '1952年3月20日，关于在' + LQ + '三反' + RQ + '运动中党员犯有贪污、浪费、官僚主义错误给予党内处分的规定.docx',
    'footnotes': [
        (LQ + '三反' + RQ + '运动', LQ + '三反' + RQ + '运动，指1951年底至1952年10月在中国党政机关、国营企业等单位开展的' + LQ + '反贪污、反浪费、反官僚主义' + RQ + '运动。'),
        ('中央节约检查委员会', '中央节约检查委员会，1951年12月成立，薄一波任主任，负责领导全国的增产节约和' + LQ + '三反' + RQ + LQ + '五反' + RQ + '运动。'),
        ('八项标准', '党员八项标准，即1951年3月第一次全国组织工作会议通过的《共产党员标准的八项条件》，是整党运动中衡量党员是否合格的基本依据。'),
        ('留党察看', '留党察看，中国共产党纪律处分之一，期限为一年或二年。受留党察看处分的党员在察看期内没有表决权、选举权和被选举权。'),
        ('支部大会', '支部大会，即党支部党员大会，是党的基层组织（支部）的最高领导机关，一般每三个月召开一次。'),
        ('党刊', '指中国共产党各级组织主办的党内刊物，用于传达中央指示和交流工作经验。'),
        ('志愿军', '中国人民志愿军，1950年10月组成开赴朝鲜参加抗美援朝战争。'),
    ],
    'end_notes': [
        '【需核】《中央节约检查委员会关于处理贪污、浪费及克服官僚主义错误的若干规定》的具体发布日期和内容需要核实。',
        '【需核】文中' + LQ + '地委书记、专员须报中央批准' + RQ + '的说法在当时组织制度中的准确性需要核查。',
    ],
}

# --- File 02: 致斯大林电（国防和经济建设）---
ANNOTATIONS['02'] = {
    'filename': '1952年3月28日，关于国防和经济建设等问题致斯大林电.docx',
    'footnotes': [
        ('朝鲜谈判', '朝鲜停战谈判，指1951年7月开始的朝鲜战争交战双方为结束战争而进行的谈判。至1952年3月，除战俘遣返问题外，其他议程已基本达成协议。'),
        ('旅顺港海军基地', '旅顺港，位于辽东半岛南端，中国北方天然良港和重要海军基地。1898年俄国租借旅顺，1905年后被日本占领。1945年苏军进驻旅顺港，根据1950年中苏条约，苏军应于1952年底前撤离。'),
        ('美日和平条约', '即《旧金山对日和约》，1951年9月8日由美国等48国与日本在美国旧金山签订。苏联、波兰、捷克斯洛伐克三国拒绝签字。中华人民共和国被排除在和约之外。'),
        ('美日安全条约', '即《日美安全保障条约》，1951年9月8日与《旧金山对日和约》同时签订，规定美国在日本及周围地区驻扎军队的权利。中方认为该条约严重威胁中国安全。'),
        ('蒙古人民共和国', '蒙古人民共和国，1924年成立，首都乌兰巴托。1949年10月16日与中华人民共和国建立外交关系。'),
        ('绥远省', '绥远省，中国旧省名，1928年设立，辖今内蒙古自治区中部地区，省会在归绥（今呼和浩特）。1954年撤销，并入内蒙古自治区。'),
        ('乌兰巴托', '乌兰巴托，蒙古人民共和国首都，原名库伦，1924年改名为乌兰巴托（意为' + LQ + '红色英雄' + RQ + '）。'),
        ('集宁', '集宁，今内蒙古自治区乌兰察布市辖区，中蒙铁路中国段的北方枢纽站。'),
        ('霍尔果斯', '霍尔果斯，位于新疆伊犁哈萨克自治州的中哈（苏）边境口岸，是中国西部重要陆路口岸。'),
        ('土耳其斯坦-西伯利亚干线', '即土西铁路（Турксиб），苏联中亚地区连接西伯利亚的铁路干线，1930年建成通车。'),
        ('萧劲光', '萧劲光（1903\u20141989），湖南长沙人，时任中国人民解放军海军司令员。'),
        ('60个步兵师', '指中苏两国关于苏联为中国陆军60个步兵师提供武器装备的军事援助计划。该计划在1951年开始执行，至1954年完成。'),
        ('橡胶树', '天然橡胶是重要的战略物资。新中国成立初期，西方国家对华实行禁运，天然橡胶供应严重不足。中国在海南岛、广东、广西、云南等地试种橡胶树以突破封锁。'),
        ('中苏股份公司', '即中苏合资股份公司，是1950年代初中苏两国在新疆有色金属、石油、民航等领域设立的合资企业。1954年赫鲁晓夫访华后苏方将股份移交中国。'),
    ],
    'end_notes': [
        '【需核】' + LQ + '哈顺' + RQ + '是否为中蒙边境实际地名，以及中蒙铁路的具体走向方案，需要核查相关资料。',
        '【需核】中苏关于旅顺港苏军撤出延期的谈判过程和时间节点，建议核对《中苏关系史》等相关资料。',
        '【需核】' + LQ + '18亿卢布' + RQ + '和' + LQ + '4亿多卢布' + RQ + '等军事贷款数额的准确性，建议核对相关外交档案。',
    ],
}

# --- File 03: 与李讷等人的谈话 ---
ANNOTATIONS['03'] = {
    'filename': '1952年3月，与李讷等人的谈话.docx',
    'footnotes': [
        ('李讷', '李讷（1940\u2014），毛泽东与江青之女。1940年生于延安，当时在北京读书。'),
        ('翟作军', '翟作军，毛泽东的警卫员（卫士），曾在毛泽东身边工作。'),
        ('江青', '江青（1914\u20141991），山东诸城人，原名李云鹤，艺名蓝苹。1938年与毛泽东结婚。时任中共中央宣传部文艺处副处长。'),
        ('毛远新', '毛远新（1941\u2014），毛泽东之弟毛泽民与朱旦华之子。毛泽民1943年在新疆被军阀盛世才杀害后，毛远新由母亲抚养。1951年回到北京。'),
        ('李敏', '李敏（1936\u2014），毛泽东与贺子珍之女，当时在北京读书。'),
    ],
    'end_notes': [
        '【需核】翟作军在毛泽东身边工作的具体时段和离开后的去向需要核查。',
        '【需核】毛远新回到北京的具体时间和背景需要核实。',
    ],
}

# --- File 04: 与菜地老农等人的谈话 ---
ANNOTATIONS['04'] = {
    'filename': '1952年3月，与菜地老农等人的谈话.docx',
    'footnotes': [
        ('王鹤滨', '王鹤滨（1924\u2014），毛泽东的保健医生，1949年至1954年任毛泽东秘书兼保健医生。'),
        ('汪东兴', '汪东兴（1916\u20142015），江西弋阳人，1932年加入中国共产党。时任中共中央办公厅警卫处处长，负责毛泽东的警卫工作。'),
        ('侯波', '侯波（1924\u20142017），女，山西夏县人，著名摄影家。时任中南海摄影科科长，毛泽东的专职摄影师。'),
        ('叶子龙', '叶子龙（1916\u20142003），湖南浏阳人，1932年参加中国工农红军。长期担任毛泽东的机要秘书，负责毛泽东的日常文书和摄影工作。'),
        ('德胜门', '德胜门，北京内城九门之一，位于北城墙西段（今德胜门立交桥附近）。明清时期军队出征多从此门出城，取' + LQ + '得胜' + RQ + '之意。现存箭楼为北京市文物保护单位。'),
        ('中南海', '中南海，位于北京故宫西侧，中海和南海的合称。1949年后为中共中央和国务院办公所在地，也是毛泽东的住所。'),
        ('周西林', '周西林，毛泽东的汽车司机。'),
        ('王振海', '王振海，毛泽东的卫士（警卫人员），负责毛泽东外出时的安全保卫工作。'),
        ('监委会', LQ + '三反' + RQ + '运动期间在基层单位成立的群众性监督组织——节约检查委员会（或称监督委员会），负责监督干部作风，受理群众检举揭发。'),
        ('军阀混战', '指1916年袁世凯死后至1928年东北易帜期间，中国各派系军阀为争夺地盘和中央政权而进行的长年混战。'),
        ('德胜门外区政府', '德胜门外区政府，1952年时北京市德胜门外地区的区级政府派出机构。当时德胜门外属北京市西四区（1952年9月更名为西单区），1958年并入西城区。' + LQ + '德胜门外区政府' + RQ + '应为西四区人民政府在德胜门外的办事处或区公所。'),
    ],
    'end_notes': [
        '【需核】菜地主人' + LQ + '姓吴' + RQ + '的身份是否在其他回忆录中有记载，建议核查。',
        '【需核】王振海在毛泽东身边工作的具体时段需要核查。',
    ],
}

# --- File 05: 荆江分洪工程题词 ---
ANNOTATIONS['05'] = {
    'filename': '1952年3月，给荆江分洪工程锦旗上的题词.docx',
    'footnotes': [
        ('荆江分洪工程', '荆江分洪工程，1952年4月5日动工、同年6月20日竣工的大型水利工程。位于湖北省荆州市长江荆江段南岸，是新中国成立后兴建的第一个大型水利枢纽工程。主体工程包括进洪闸（北闸）、节制闸（南闸）和荆江分洪区围堤等，主要目的是减轻荆江大堤的洪水威胁，保护江汉平原和武汉三镇的安全。毛泽东、周恩来亲自批准工程方案。'),
    ],
    'end_notes': [],
}


# ═════════════════ MAIN ═════════════════
def main():
    all_files = [f for f in os.listdir(MONTH_DIR) if f.endswith('.docx') and not f.startswith('__')]
    for key, data in ANNOTATIONS.items():
        fname = data['filename']
        matched = None
        for af in all_files:
            if af.startswith(fname[:15]):
                matched = af
                break
        if matched:
            data['filename'] = matched
            process_file(data)
        else:
            print('NOT FOUND: ' + fname[:30])
    for f in os.listdir(MONTH_DIR):
        if f.startswith('__'):
            os.remove(os.path.join(MONTH_DIR, f))
    print('\n3月初注完成!')


if __name__ == '__main__':
    main()
