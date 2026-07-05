#!/usr/bin/env python3
"""7月份批量初注脚本"""
import os, shutil
from docx import Document
from lxml import etree

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W14NS = 'http://schemas.microsoft.com/office/word/2010/wordml'
MONTH_DIR = r'E:\1952年大传\1952年大传\1952.7'
OUT_PREFIX = '_LLM初注'
LQ='\u201c'; RQ='\u201d'

FOOTNOTE_TPL = (
    '<w:footnote xmlns:w="%s" xmlns:w14="%s" w:id="{id}">'
    '<w:p><w:pPr><w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr></w:pPr>'
    '<w:r><w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr><w:footnoteRef/></w:r>'
    '<w:r><w:rPr>'
    '<w:rFonts w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" '
    'w:eastAsiaTheme="minorEastAsia" w:cstheme="minorBidi"/>'
    '<w:kern w:val="2"/><w:sz w:val="21"/><w:szCs w:val="21"/>'
    '<w:lang w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'
    '</w:rPr><w:t xml:space="preserve">{text}</w:t></w:r>'
    '</w:p></w:footnote>'
) % (WNS, W14NS)

def esc(s): return s.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')

def get_max_id(fn_root): return max(int(fn.get('{%s}id'%WNS,'0')) for fn in fn_root.findall('{%s}footnote'%WNS))

def insert_after_text(paragraph, search_text, fn_id):
    full_text = paragraph.text
    if search_text not in full_text: return False
    end_pos = full_text.index(search_text) + len(search_text)
    cur = 0
    for run in paragraph.runs:
        run_end = cur + len(run.text)
        if cur <= end_pos <= run_end:
            sp = end_pos - cur; after = run.text[sp:]; run.text = run.text[:sp]
            ref = etree.fromstring(('<w:r xmlns:w="%s"><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr><w:footnoteReference w:id="%d"/></w:r>')%(WNS,fn_id))
            run._element.addnext(ref)
            if after:
                a = etree.fromstring(('<w:r xmlns:w="%s"><w:rPr></w:rPr><w:t xml:space="preserve">%s</w:t></w:r>')%(WNS,after))
                ref.addnext(a)
            return True
        cur = run_end
    return False

def add_fn(fn_part, fn_id, fn_text):
    r = etree.fromstring(fn_part.blob)
    r.append(etree.fromstring(FOOTNOTE_TPL.format(id=fn_id, text=esc(fn_text))))
    fn_part._blob = etree.tostring(r, xml_declaration=True, encoding='UTF-8', standalone=True)

def add_end(doc, items):
    if not items: return
    doc.add_paragraph().add_run('\n' + '\u2500' * 40)
    r = doc.add_paragraph().add_run('本篇待人工确认事项：'); r.bold = True
    for i, item in enumerate(items, 1): doc.add_paragraph().add_run('%d. %s' % (i, item))

def process_file(data):
    fname = data['filename']; src = os.path.join(MONTH_DIR, fname)
    dst = os.path.join(MONTH_DIR, fname.replace('.docx','_'+OUT_PREFIX+'.docx'))
    if os.path.exists(dst): print('SKIP'); return
    shutil.copy2(src, dst); doc = Document(dst)
    fn_part = None
    for rel in doc.part.rels.values():
        if 'footnote' in rel.reltype.lower(): fn_part = rel.target_part; break
    if fn_part is None: print('ERROR: no fn'); return
    next_id = get_max_id(etree.fromstring(fn_part.blob)) + 1
    print('>>> ' + fname[:30] + ' (id=%d)' % next_id)
    ins = 0
    for st, ft in data['footnotes']:
        for para in doc.paragraphs:
            if st in para.text:
                if insert_after_text(para, st, next_id): add_fn(fn_part, next_id, ft); next_id += 1; ins += 1; break
        else: print('  NF: ' + st[:50])
    if data.get('end_notes'): add_end(doc, data['end_notes'])
    doc.save(dst); print('  -> %d' % ins)

# ═══════ ANNOTATIONS ═══════
ANNOTATIONS = {}
S3=LQ+'三反'+RQ; S5=LQ+'五反'+RQ

# 01: 致斯大林电 (7.18)
ANNOTATIONS['01'] = {
    'filename': '1952年7月18日，致斯大林电.docx',
    'footnotes': [
        ('金日成', '金日成（1912\u20141994），原名金成柱，朝鲜民主主义人民共和国的创建者。1948年起任朝鲜民主主义人民共和国内阁首相，1949年起任朝鲜劳动党中央委员会委员长。抗美援朝期间与毛泽东、斯大林保持频繁电文往来，协调战争指导方针。'),
        ('李克农', '李克农（1899\u20141962），安徽巢县人，1926年加入中国共产党。1950年任外交部副部长。抗美援朝战争期间，受中央委派在幕后领导和指导朝鲜停战谈判的中方代表团。1955年被授予上将军衔。'),
        ('三八线', '三八线，即北纬38度线。1945年8月日本投降后，美苏两国商定以此线为界分别接受朝鲜半岛日军的投降，此后成为朝鲜和韩国事实上的分界线。朝鲜战争爆发后，战线多次反复穿越三八线，最终停战协定基本以三八线附近的实际接触线为军事分界线。'),
        ('长津湖发电站', '长津湖发电站，位于朝鲜咸镜南道长津郡，是朝鲜北部重要的水力发电站之一。长津湖地区在1950年11月至12月爆发了著名的长津湖战役。'),
        ('水丰发电站', '水丰发电站，位于鸭绿江中游水丰水库（中朝界河），是中朝边界上最大的水力发电站，1940年建成。1952年6月23日遭美军飞机集中轰炸严重受损。'),
        ('北仓发电站', '北仓发电站，位于朝鲜平安南道北仓郡，是朝鲜北部的主要火力发电厂之一。'),
        ('中国人民志愿军', '中国人民志愿军，1950年10月组成并开赴朝鲜参加抗美援朝战争。彭德怀任司令员兼政治委员。'),
    ],
    'end_notes': [
        '【需核】毛泽东7月15日3时致金日成电原文和斯大林（菲利波夫）对毛泽东7月18日电的复电内容，建议核查相关外交档案。',
        '【需核】金日成7月16日复电中提到的' + LQ + '高射炮部队至少10个团' + RQ + '等具体要求是否得到满足，建议核实后续军援情况。',
    ],
}

# 02: 炮兵弹药致斯大林电
ANNOTATIONS['02'] = {
    'filename': '1952年7月4日，关于补充炮兵弹药和器材问题致斯大林电.docx',
    'footnotes': [
        ('水丰水力发电站', '水丰发电站，位于鸭绿江中游水丰水库（中朝界河），是中朝边界上最大的水力发电站。1952年6月23日遭美军飞机集中轰炸严重受损。'),
        ('鸭绿江', '鸭绿江，发源于长白山，流经中国吉林、辽宁两省和朝鲜两江道、慈江道、平安北道，注入黄海，全长795公里，是中朝两国的界河。'),
        ('军事贷款', '军事贷款，指苏联向中国提供的用于购买武器装备和军用物资的低息贷款。中苏两国在1951年2月1日签订了军事贷款协定。'),
    ],
    'end_notes': [
        '【需核】1952年5月31日毛泽东致斯大林电附录中的炮兵武器弹药申请单的具体内容需要核查。',
        '【需核】斯大林6月13日来电（提及朝鲜停战谈判中敌方' + LQ + '似乎会有转变' + RQ + '）的内容需要核实。',
    ],
}

# 03: 增产节约指示
ANNOTATIONS['03'] = {
    'filename': '1952年7月8日，关于目前开展增产节约运动中应注意的问题的指示.docx',
    'footnotes': [
        ('增产节约', '增产节约运动，新中国成立初期在全国范围内开展的以增加生产、厉行节约为中心内容的群众运动。1951年10月毛泽东向全国发出开展增产节约运动的号召。'),
        (S3, S3 + '运动，指1951年底至1952年10月在中国党政机关、国营企业等单位开展的' + LQ + '反贪污、反浪费、反官僚主义' + RQ + '运动。'),
        ('经济核算制', '经济核算制，社会主义国营企业管理的基本制度之一，要求企业以货币形式独立核算生产经营的消耗和成果，以收抵支，保证盈利。'),
        (S5, S5 + '运动，指1952年上半年在资本主义工商业者中开展的' + LQ + '反行贿、反偷税漏税、反盗骗国家财产、反偷工减料、反盗窃国家经济情报' + RQ + '的运动。'),
        (LQ + '五毒' + RQ, LQ + '五毒' + RQ + '，指' + S5 + '运动中所反对的五种违法行为：行贿、偷税漏税、盗骗国家财产、偷工减料、盗窃国家经济情报。'),
        ('劳资两利', LQ + '劳资两利' + RQ + '，1949年4月毛泽东在《在中国共产党第七届中央委员会第二次全体会议上的报告》中提出' + LQ + '发展生产、繁荣经济、公私兼顾、劳资两利' + RQ + '的新民主主义经济方针。' + LQ + '劳资两利' + RQ + '要求兼顾工人的合理权益和资本家的合法利润。'),
        ('集体合同', '集体合同，指工会代表职工与企业（资方）就劳动报酬、工作时间、劳动条件等事项签订的书面协议。新中国成立初期是处理劳资关系的重要法律形式。'),
    ],
    'end_notes': [
        '【需核】华东工业会议的报告和东北、华北关于增产节约运动的报告的具体内容需要核查。',
    ],
}

def main():
    all_files = [f for f in os.listdir(MONTH_DIR) if f.endswith('.docx') and not f.startswith('__')]
    for key, data in ANNOTATIONS.items():
        fname = data['filename']; matched = None
        for af in all_files:
            if af.startswith(fname[:15]): matched = af; break
        if matched: data['filename'] = matched; process_file(data)
        else: print('NF: ' + fname[:30])
    for f in os.listdir(MONTH_DIR):
        if f.startswith('__'): os.remove(os.path.join(MONTH_DIR, f))
    print('\n7月初注完成!')

if __name__ == '__main__': main()
